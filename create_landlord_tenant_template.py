from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_landlord_tenant_ontario_template():
    """Create a specialized template for Ontario landlord/tenant disputes"""
    document = Document()
    
    # Set up document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add heading
    heading = document.add_heading('LANDLORD/TENANT DISPUTE LETTER', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add recipient line
    document.add_paragraph("To Whom It May Concern,")
    document.add_paragraph()  # Blank line
    
    # Add introduction
    intro = document.add_paragraph()
    intro.add_run("My name is {{ name }}, and I am a resident of the province of {{ province }}.")
    document.add_paragraph("I am writing to formally raise a concern regarding my tenancy. The issue I am experiencing is as follows:")
    
    # Add description placeholder
    document.add_paragraph()  # Blank line
    description_para = document.add_paragraph()
    description_para.add_run("{{ description }}")
    document.add_paragraph()  # Blank line
    
    # Add legal reference
    legal_para = document.add_paragraph()
    legal_para.add_run("This dispute falls under the jurisdiction of the Residential Tenancies Act, 2006 (Ontario). I am requesting that this matter be addressed promptly and in accordance with the rights afforded to tenants under provincial legislation.")
    
    # Add documentation note
    document.add_paragraph()  # Blank line
    document.add_paragraph("Please find all necessary documentation attached to support my claim.")
    
    # Add signature block
    document.add_paragraph()  # Blank line
    document.add_paragraph("Sincerely,")
    document.add_paragraph("{{ name }}")
    document.add_paragraph("{{ email }}")
    
    # Add date
    document.add_paragraph()  # Blank line
    date_para = document.add_paragraph()
    date_para.add_run("Date: {{ now() }}")
    
    # Ensure templates directory exists
    os.makedirs("templates/docs", exist_ok=True)
    
    # Save the document
    filename = "templates/docs/landlord_tenant_ON.docx"
    document.save(filename)
    print(f"Created template: {filename}")
    
    return filename

if __name__ == "__main__":
    template_file = create_landlord_tenant_ontario_template()
    print(f"Successfully created template at {template_file}")