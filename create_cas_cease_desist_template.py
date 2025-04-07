from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def get_cas_provincial_legislation(province_code):
    """Return the appropriate child protection legislation name based on province code"""
    legislation_map = {
        "ON": "Child, Youth and Family Services Act",
        "BC": "Child, Family and Community Service Act",
        "AB": "Child, Youth and Family Enhancement Act",
        "QC": "Youth Protection Act (Loi sur la protection de la jeunesse)",
        "MB": "Child and Family Services Act",
        "SK": "Child and Family Services Act",
        "NS": "Children and Family Services Act",
        "NB": "Family Services Act",
        "NL": "Children, Youth and Families Act",
        "PE": "Child Protection Act",
        "NT": "Child and Family Services Act",
        "NU": "Child and Family Services Act",
        "YT": "Child and Family Services Act"
    }
    
    return legislation_map.get(province_code, "applicable child protection legislation")

def get_agency_name(province_code):
    """Return the appropriate agency name based on province code"""
    agency_map = {
        "ON": "Children's Aid Society",
        "BC": "Ministry of Children and Family Development",
        "AB": "Children's Services",
        "QC": "Direction de la protection de la jeunesse (DPJ)",
        "MB": "Child and Family Services",
        "SK": "Ministry of Social Services",
        "NS": "Department of Community Services",
        "NB": "Department of Social Development",
        "NL": "Department of Children, Seniors and Social Development",
        "PE": "Child Protection Services",
        "NT": "Child and Family Services",
        "NU": "Department of Family Services",
        "YT": "Family and Children's Services"
    }
    
    return agency_map.get(province_code, "Child Protection Agency")

def create_french_quebec_template():
    """Create a specialized French template for Quebec DPJ"""
    document = Document()
    
    # Set up document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add sender info placeholders
    sender_info = document.add_paragraph()
    sender_info.add_run("{{ name }}\n")
    sender_info.add_run("{{ address }}\n")
    sender_info.add_run("{{ email }}")
    sender_info.add_run("\nQC")
    sender_info.add_run("\n{{ now() }}")
    
    # Add space after sender info
    document.add_paragraph()
    
    # Add recipient info placeholders
    recipient_info = document.add_paragraph()
    recipient_info.add_run("{{ recipient_name }}\n")
    recipient_info.add_run("{{ recipient_address }}")
    
    # Add space after recipient info
    document.add_paragraph()
    
    # Add subject line
    subject = document.add_paragraph()
    subject_run = subject.add_run("Objet: Mise en demeure formelle – Harcèlement et ingérence illégale")
    subject_run.bold = True
    
    # Add greeting
    document.add_paragraph("À qui de droit,")
    
    # Add main body - French version
    main_body = document.add_paragraph()
    main_body.add_run("Cette lettre constitue une mise en demeure formelle de cesser immédiatement toute ingérence injustifiée, excessive ou illégale dans ma famille et mon foyer par vous ou tout représentant de la Direction de la protection de la jeunesse (DPJ).")
    
    document.add_paragraph()
    legal_para = document.add_paragraph()
    legal_para.add_run("Bien que je reconnaisse le mandat de votre organisme en vertu de la ")
    legal_para.add_run("Loi sur la protection de la jeunesse").bold = True
    legal_para.add_run(" au Québec, vos actions récentes semblent dépasser la portée de l'autorité légale. Cela inclut, sans s'y limiter:")
    
    # Add list of actions - French
    document.add_paragraph()
    bullet_points = document.add_paragraph("- Visites non annoncées ou répétitives sans justification\n")
    bullet_points.add_run("- Communication intimidante\n")
    bullet_points.add_run("- Actions constituant du harcèlement, de la diffamation ou un abus de pouvoir\n")
    bullet_points.add_run("- Tentatives de miner mon autorité ou mes droits parentaux sans base légale")
    
    # Add demand - French
    document.add_paragraph()
    document.add_paragraph("J'exige que vous cessiez immédiatement ces actions. Toute ingérence ou harcèlement continu pourrait entraîner le dépôt d'une plainte formelle auprès de votre organisme de surveillance provincial, ainsi qu'une action en justice pour dommages, discrimination ou violation des droits civils et parentaux.")
    
    # Add notice - French
    document.add_paragraph()
    document.add_paragraph("Que cette lettre serve d'avis que je m'attends à ce que toutes les communications futures soient menées par écrit, à moins qu'une urgence ne le justifie autrement. Il vous est également conseillé de conserver toutes les notes internes, registres et correspondances liés à cette affaire comme preuves potentielles.")
    
    # Add signature
    document.add_paragraph()
    signature = document.add_paragraph("Sincèrement,")
    signature.add_run("\n{{ name }}")
    signature.add_run("\n{{ email }}")
    
    # Ensure templates directory exists
    os.makedirs("templates/docs", exist_ok=True)
    
    # Save the document
    filename = "templates/docs/cas_cease_desist_QC_FR.docx"
    document.save(filename)
    print(f"Created template: {filename}")
    
    return filename

def create_cas_cease_desist_template(province_code=None):
    """Create a specialized template for CAS cease and desist letters"""
    document = Document()
    
    # Set up document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Get province-specific legislation and agency name if provided
    legal_act = get_cas_provincial_legislation(province_code) if province_code else "{{ legal_act }}"
    province = province_code if province_code else "{{ province }}"
    agency_name = get_agency_name(province_code) if province_code else "{{ agency_name }}"
    
    # Add sender info placeholders
    sender_info = document.add_paragraph()
    sender_info.add_run("{{ name }}\n")
    sender_info.add_run("{{ address }}\n")
    sender_info.add_run("{{ email }}")
    if province_code:
        sender_info.add_run(f"\n{province}")
    sender_info.add_run("\n{{ now() }}")
    
    # Add space after sender info
    document.add_paragraph()
    
    # Add recipient info placeholders
    recipient_info = document.add_paragraph()
    recipient_info.add_run("{{ recipient_name }}\n")
    recipient_info.add_run("{{ recipient_address }}")
    
    # Add space after recipient info
    document.add_paragraph()
    
    # Add subject line
    subject = document.add_paragraph()
    subject_run = subject.add_run("Re: Formal Cease and Desist Demand – Harassment and Unlawful Interference")
    subject_run.bold = True
    
    # Add greeting
    document.add_paragraph("To Whom It May Concern,")
    
    # Add main body
    main_body = document.add_paragraph()
    main_body.add_run("This letter serves as a formal demand to ")
    cease_run = main_body.add_run("cease and desist all unwarranted, excessive, or unlawful interference")
    cease_run.bold = True
    main_body.add_run(f" with my family and household by you or any representatives of {agency_name}.")
    
    # Add legal reference
    document.add_paragraph()
    legal_para = document.add_paragraph()
    legal_para.add_run(f"While I recognize the mandate of your agency under the ")
    legal_act_run = legal_para.add_run(f"{legal_act}")
    legal_act_run.bold = True
    legal_para.add_run(f" in ")
    province_run = legal_para.add_run(f"{province}")
    province_run.bold = True
    legal_para.add_run(", your recent actions appear to exceed the scope of lawful authority. This includes, but is not limited to:")
    
    # Add list of actions
    document.add_paragraph()
    bullet_points = document.add_paragraph("- Unannounced or repetitive visits without justification\n")
    bullet_points.add_run("- Intimidating communication\n")
    bullet_points.add_run("- Actions that amount to harassment, defamation, or abuse of power\n")
    bullet_points.add_run("- Attempts to undermine my parental authority or rights without legal basis")
    
    # Add demand
    document.add_paragraph()
    document.add_paragraph("I demand that you immediately cease these actions. Continued overreach or harassment may result in a formal complaint being filed with your provincial oversight body, as well as legal action for damages, discrimination, or violation of civil and parental rights.")
    
    # Add notice
    document.add_paragraph()
    document.add_paragraph("Let this letter serve as notice that I expect all further communications to be conducted in writing, unless an emergency justifies otherwise. You are also advised to preserve all internal notes, records, and correspondence related to this matter as potential evidence.")
    
    # Add signature
    document.add_paragraph()
    signature = document.add_paragraph("Sincerely,")
    signature.add_run("\n{{ name }}")
    signature.add_run("\n{{ email }}")
    
    # Ensure templates directory exists
    os.makedirs("templates/docs", exist_ok=True)
    
    # Save the document with province-specific filename if provided
    if province_code:
        filename = f"templates/docs/cas_cease_desist_{province_code}.docx"
    else:
        filename = "templates/docs/cas_cease_desist.docx"
    
    document.save(filename)
    print(f"Created template: {filename}")
    
    return filename

def create_all_provincial_templates():
    """Create CAS cease and desist templates for all provinces"""
    provinces = ["ON", "BC", "AB", "MB", "SK", "NS", "NB", "NL", "PE", "NT", "NU", "YT"]
    created_files = []
    
    for province in provinces:
        filename = create_cas_cease_desist_template(province)
        created_files.append(filename)
    
    # Add Quebec French version
    filename = create_french_quebec_template()
    
    # Add Quebec English version
    filename = create_cas_cease_desist_template("QC")
    created_files.append(filename)
    
    return created_files

if __name__ == "__main__":
    # Create generic template
    generic_template = create_cas_cease_desist_template()
    print(f"Created generic template: {generic_template}")
    
    # Create province-specific templates
    provincial_templates = create_all_provincial_templates()
    print(f"Created {len(provincial_templates)} provincial templates.")