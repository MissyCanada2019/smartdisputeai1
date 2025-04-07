#!/usr/bin/env python3
"""
Template Creation Script for CAS Answer and Plan of Care Documents
Creates templates for all Canadian provinces with the appropriate legal references
"""

import os
import sys
from docxtpl import DocxTemplate
from docx import Document
import docx

def get_cas_provincial_legislation(province_code):
    """Return the appropriate child protection legislation name based on province code"""
    legislation_map = {
        "ON": "Child, Youth and Family Services Act",
        "BC": "Child, Family and Community Service Act",
        "AB": "Child, Youth and Family Enhancement Act",
        "QC": "Youth Protection Act",
        "MB": "Child and Family Services Act",
        "SK": "Child and Family Services Act",
        "NS": "Children and Family Services Act",
        "NB": "Family Services Act",
        "NL": "Children, Youth and Families Act",
        "PE": "Child Protection Act",
        "NT": "Child and Family Services Act",
        "NU": "Child and Family Services Act",
        "YT": "Child and Family Services Act"
    }
    
    return legislation_map.get(province_code, "Child Protection Legislation")

def get_agency_name(province_code):
    """Return the appropriate agency name based on province code"""
    agency_map = {
        "ON": "Children's Aid Society",
        "BC": "Ministry of Children and Family Development",
        "AB": "Child and Family Services",
        "QC": "Direction de la protection de la jeunesse (DPJ)",
        "MB": "Child and Family Services",
        "SK": "Ministry of Social Services",
        "NS": "Department of Community Services",
        "NB": "Department of Social Development",
        "NL": "Department of Children, Seniors and Social Development",
        "PE": "Child Protection Services",
        "NT": "Child and Family Services",
        "NU": "Department of Family Services",
        "YT": "Family and Children's Services"
    }
    
    return agency_map.get(province_code, "Child Protection Services")

def get_plan_of_care_details(province_code):
    """Return province-specific plan of care terminology and requirements"""
    details_map = {
        "ON": "Plan of Care as defined in section 109 of the Child, Youth and Family Services Act",
        "BC": "Care Plan as required under section 35 of the Child, Family and Community Service Act",
        "AB": "Service Plan as outlined in section 9 of the Child, Youth and Family Enhancement Act",
        "QC": "Intervention Plan as required by section 32 of the Youth Protection Act",
        "MB": "Service Plan as defined in section 13 of the Child and Family Services Act",
        "SK": "Case Plan as outlined in section 8 of the Child and Family Services Act",
        "NS": "Plan of Care as defined in section 13 of the Children and Family Services Act",
        "NB": "Child Service Plan as required by the Family Services Act",
        "NL": "Plan for the Child as outlined in the Children, Youth and Families Act",
        "PE": "Service Plan pursuant to section 15 of the Child Protection Act",
        "NT": "Plan of Care as defined in the Child and Family Services Act",
        "NU": "Plan of Care agreement under the Child and Family Services Act",
        "YT": "Case Plan as required by the Child and Family Services Act"
    }
    
    return details_map.get(province_code, "Plan of Care as required by provincial child protection legislation")

def create_french_quebec_template():
    """Create a specialized French template for Quebec DPJ"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    output_path = f"{template_dir}/cas_answer_plan_QC_FR.docx"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('RÉPONSE ET PLAN D\'INTERVENTION', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add DPJ info
    doc.add_paragraph("À:")
    doc.add_paragraph("Direction de la protection de la jeunesse (DPJ)")
    doc.add_paragraph("{{agency_address}}")
    doc.add_paragraph("{{agency_city}}, QC {{agency_postal_code}}")
    doc.add_paragraph()
    
    # Add recipient info
    doc.add_paragraph("De:")
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("{{parent_address}}")
    doc.add_paragraph("{{parent_city}}, QC {{parent_postal_code}}")
    doc.add_paragraph()
    
    # Add reference information
    doc.add_paragraph("Référence: {{case_reference}}")
    doc.add_paragraph("Intervenant(e): {{caseworker_name}}")
    doc.add_paragraph("Concernant: {{child_name}}, né(e) le {{child_dob}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph()
    
    doc.add_paragraph("Je vous écris en réponse aux allégations et préoccupations soulevées dans votre communication du {{allegation_date}}, conformément à la Loi sur la protection de la jeunesse.")
    doc.add_paragraph()
    
    # Response to allegations
    doc.add_paragraph("RÉPONSE AUX ALLÉGATIONS:", style='Heading2')
    doc.add_paragraph("{{allegations_response}}")
    doc.add_paragraph()
    
    # Proposed plan of care
    doc.add_paragraph("PLAN D'INTERVENTION PROPOSÉ:", style='Heading2')
    doc.add_paragraph("Conformément aux exigences de l'article 32 de la Loi sur la protection de la jeunesse concernant le Plan d'intervention, je propose les mesures suivantes:")
    doc.add_paragraph("{{proposed_plan}}")
    doc.add_paragraph()
    
    # Additional resources and support
    doc.add_paragraph("RESSOURCES ET SOUTIEN SUPPLÉMENTAIRES:", style='Heading2')
    doc.add_paragraph("{{additional_resources}}")
    doc.add_paragraph()
    
    # Request for meeting
    doc.add_paragraph("Je souhaite planifier une réunion pour discuter de ce plan d'intervention et travailler en collaboration pour assurer le bien-être de {{child_name}}. Je suis disponible aux dates et heures suivantes:")
    doc.add_paragraph("{{available_dates}}")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Je vous remercie de votre attention à cette affaire. Je suis engagé(e) à travailler de manière collaborative pour assurer le meilleur intérêt de mon enfant, tout en respectant mes droits parentaux en vertu de la Loi sur la protection de la jeunesse.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Veuillez agréer mes salutations distinguées,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("Téléphone: {{parent_phone}}")
    doc.add_paragraph("Courriel: {{parent_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_cas_answer_plan_template(province_code=None):
    """Create a specialized template for CAS Answer and Plan of Care"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    # Create a base template name
    base_template_path = f"{template_dir}/cas_answer_plan.docx"
    
    # If a province code is provided, create a province-specific template
    if province_code:
        output_path = f"{template_dir}/cas_answer_plan_{province_code}.docx"
        legislation = get_cas_provincial_legislation(province_code)
        agency_name = get_agency_name(province_code)
        plan_details = get_plan_of_care_details(province_code)
    else:
        output_path = base_template_path
        legislation = "Child Protection Act"
        agency_name = "Child Protection Services"
        plan_details = "Plan of Care as required by provincial child protection legislation"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('RESPONSE TO ALLEGATIONS AND PROPOSED PLAN OF CARE', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add agency info
    doc.add_paragraph("To:")
    doc.add_paragraph(f"{agency_name}")
    doc.add_paragraph("{{agency_address}}")
    doc.add_paragraph("{{agency_city}}, {{agency_province}} {{agency_postal_code}}")
    doc.add_paragraph()
    
    # Add parent info
    doc.add_paragraph("From:")
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("{{parent_address}}")
    doc.add_paragraph("{{parent_city}}, {{parent_province}} {{parent_postal_code}}")
    doc.add_paragraph()
    
    # Add reference information
    doc.add_paragraph("Reference: {{case_reference}}")
    doc.add_paragraph("Caseworker: {{caseworker_name}}")
    doc.add_paragraph("Regarding: {{child_name}}, born {{child_dob}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph("Dear Sir/Madam,")
    doc.add_paragraph()
    
    doc.add_paragraph(f"I am writing in response to the allegations and concerns raised in your communication dated {{allegation_date}}, in accordance with the {legislation}.")
    doc.add_paragraph()
    
    # Response to allegations
    doc.add_paragraph("RESPONSE TO ALLEGATIONS:", style='Heading2')
    doc.add_paragraph("{{allegations_response}}")
    doc.add_paragraph()
    
    # Proposed plan of care
    doc.add_paragraph("PROPOSED PLAN OF CARE:", style='Heading2')
    doc.add_paragraph(f"In accordance with the requirements for a {plan_details}, I propose the following measures:")
    doc.add_paragraph("{{proposed_plan}}")
    doc.add_paragraph()
    
    # Additional resources and support
    doc.add_paragraph("ADDITIONAL RESOURCES AND SUPPORT:", style='Heading2')
    doc.add_paragraph("{{additional_resources}}")
    doc.add_paragraph()
    
    # Request for meeting
    doc.add_paragraph("I would like to schedule a meeting to discuss this plan of care and work collaboratively to ensure the well-being of {{child_name}}. I am available on the following dates and times:")
    doc.add_paragraph("{{available_dates}}")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph(f"Thank you for your attention to this matter. I am committed to working collaboratively to ensure the best interest of my child, while respecting my parental rights under the {legislation}.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("Phone: {{parent_phone}}")
    doc.add_paragraph("Email: {{parent_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_all_provincial_templates():
    """Create CAS answer and plan templates for all provinces"""
    provinces = ["ON", "BC", "AB", "QC", "MB", "SK", "NS", "NB", "NL", "PE", "NT", "NU", "YT"]
    template_count = 0
    
    # First create generic template
    create_cas_answer_plan_template()
    
    # Then create province-specific templates
    for province in provinces:
        create_cas_answer_plan_template(province)
        template_count += 1
    
    # Create Quebec French template
    create_french_quebec_template()
    
    print(f"Created {template_count} provincial CAS answer and plan templates.")

if __name__ == "__main__":
    create_all_provincial_templates()