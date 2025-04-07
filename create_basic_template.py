from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_template(province, dispute_type):
    """Create a basic template document for a specific province and dispute type"""
    document = Document()
    
    # Set up document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add sender information placeholder
    sender_para = document.add_paragraph()
    sender_para.add_run("{{name}}").bold = True
    document.add_paragraph("{{address}}")
    document.add_paragraph("{{email}}")
    
    # Add date
    document.add_paragraph("{{date}}")
    
    # Add recipient information
    document.add_paragraph()  # Blank line
    document.add_paragraph("{{recipient_name}}")
    document.add_paragraph("{{recipient_address}}")
    
    # Add subject line
    document.add_paragraph()  # Blank line
    subject = document.add_paragraph()
    subject_run = subject.add_run("RE: ")
    subject_run.bold = True
    
    # Customize subject based on dispute type
    if dispute_type == "landlord_tenant":
        subject.add_run("Landlord/Tenant Dispute - Legal Notice").bold = True
    elif dispute_type == "credit":
        subject.add_run("Formal Credit Dispute Notice").bold = True
    elif dispute_type == "cas":
        subject.add_run("Children's Aid Society (CAS) Formal Complaint").bold = True
    
    # Add salutation
    document.add_paragraph()  # Blank line
    document.add_paragraph("Dear {{recipient_name}},")
    
    # Add opening paragraph based on dispute type and province
    document.add_paragraph()  # Blank line
    opening = document.add_paragraph()
    
    # Customize opening based on dispute type and province
    if dispute_type == "landlord_tenant":
        if province == "ON":
            opening.add_run("I am writing regarding a landlord/tenant matter under the Ontario Residential Tenancies Act, 2006. ")
        elif province == "BC":
            opening.add_run("I am writing regarding a landlord/tenant matter under the British Columbia Residential Tenancy Act. ")
        elif province == "AB":
            opening.add_run("I am writing regarding a landlord/tenant matter under the Alberta Residential Tenancies Act. ")
        elif province == "QC":
            opening.add_run("I am writing regarding a landlord/tenant matter under the Quebec Civil Code. ")
    elif dispute_type == "credit":
        opening.add_run("I am writing to dispute information on my credit report. Under the Consumer Reporting Act, I have the right to dispute inaccurate information. ")
    elif dispute_type == "cas":
        if province == "ON":
            opening.add_run("I am writing regarding a matter involving the Children's Aid Society under the Ontario Child, Youth and Family Services Act. ")
        elif province == "BC":
            opening.add_run("I am writing regarding a matter involving the Ministry of Children and Family Development under the British Columbia Child, Family and Community Service Act. ")
        elif province == "AB":
            opening.add_run("I am writing regarding a matter involving Child and Family Services under the Alberta Child, Youth and Family Enhancement Act. ")
        elif province == "QC":
            opening.add_run("I am writing regarding a matter involving the Director of Youth Protection under the Quebec Youth Protection Act. ")
    
    # Add placeholder for dispute details
    document.add_paragraph("{{description}}")
    
    # Add closing statement based on dispute type
    closing = document.add_paragraph()
    if dispute_type == "landlord_tenant":
        closing.add_run("I expect your prompt attention to this matter and a response within 14 days. If I do not receive a satisfactory response, I may pursue additional legal remedies available to me under the applicable legislation.")
    elif dispute_type == "credit":
        closing.add_run("Under the Consumer Reporting Act, you are required to investigate this matter and respond within 30 days. Please provide written confirmation that the disputed information has been corrected.")
    elif dispute_type == "cas":
        closing.add_run("I request your immediate attention to this matter. Please provide a written response to this complaint within 14 days, detailing the steps you will take to address these concerns.")
    
    # Add signature block
    document.add_paragraph()  # Blank line
    document.add_paragraph("Sincerely,")
    document.add_paragraph()  # Space for signature
    document.add_paragraph("{{name}}")
    
    # Ensure templates directory exists
    os.makedirs("templates/docs", exist_ok=True)
    
    # Save the document
    filename = f"templates/docs/{dispute_type}_{province}.docx"
    document.save(filename)
    print(f"Created template: {filename}")
    
    return filename

def create_all_templates():
    """Create templates for all province and dispute type combinations"""
    provinces = ["ON", "BC", "AB", "QC"]
    dispute_types = ["landlord_tenant", "credit", "cas"]
    
    created_files = []
    for province in provinces:
        for dispute_type in dispute_types:
            filename = create_template(province, dispute_type)
            created_files.append(filename)
    
    return created_files

if __name__ == "__main__":
    templates = create_all_templates()
    print(f"Created {len(templates)} template files.")