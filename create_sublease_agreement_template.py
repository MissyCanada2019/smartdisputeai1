#!/usr/bin/env python3
"""
Template Creation Script for Sublease Agreement Documents
Creates templates for all Canadian provinces with the appropriate legal references
"""

import os
import sys
from docxtpl import DocxTemplate
from docx import Document
import docx

def get_provincial_legislation(province_code):
    """Return the appropriate legislation name based on province code"""
    legislation_map = {
        "ON": "Residential Tenancies Act",
        "BC": "Residential Tenancy Act",
        "AB": "Residential Tenancy Act",
        "QC": "Civil Code of Quebec",
        "MB": "Residential Tenancies Act",
        "SK": "Residential Tenancies Act",
        "NS": "Residential Tenancies Act",
        "NB": "Residential Tenancies Act",
        "NL": "Residential Tenancies Act",
        "PE": "Rental of a Residential Property Act",
        "NT": "Residential Tenancies Act",
        "NU": "Residential Tenancies Act",
        "YT": "Residential Landlord and Tenant Act"
    }
    
    return legislation_map.get(province_code, "Residential Tenancies Act")

def get_sublease_restrictions(province_code):
    """Return province-specific sublease restrictions and requirements"""
    restrictions_map = {
        "ON": "requires written consent from the landlord unless the lease specifies otherwise",
        "BC": "requires written consent from the landlord, which cannot be unreasonably withheld",
        "AB": "requires written consent from the landlord, which can be reasonably withheld",
        "QC": "can be done without landlord consent, but with prior notice, unless the lease prohibits it",
        "MB": "requires written consent from the landlord, which cannot be unreasonably withheld",
        "SK": "requires written consent from the landlord, which cannot be unreasonably withheld",
        "NS": "requires written consent from the landlord, which cannot be unreasonably withheld",
        "NB": "requires written consent from the landlord",
        "NL": "requires written consent from the landlord",
        "PE": "requires written consent from the landlord",
        "NT": "requires written consent from the landlord",
        "NU": "requires written consent from the landlord",
        "YT": "requires written consent from the landlord, which cannot be unreasonably withheld"
    }
    
    return restrictions_map.get(province_code, "may require written consent from the landlord depending on your specific lease and local regulations")

def create_french_quebec_template():
    """Create a specialized French template for Quebec"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    output_path = f"{template_dir}/sublease_agreement_QC_FR.docx"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('CONTRAT DE SOUS-LOCATION', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    doc.add_paragraph()
    
    # Add introduction
    doc.add_paragraph("CE CONTRAT DE SOUS-LOCATION (« Contrat ») est conclu entre :")
    doc.add_paragraph()
    
    # Add parties
    doc.add_paragraph("LOCATAIRE PRINCIPAL (« Sous-locateur ») :")
    doc.add_paragraph("Nom : {{tenant_name}}")
    doc.add_paragraph("Adresse actuelle : {{property_address}}, {{property_city}}, QC {{property_postal_code}}")
    doc.add_paragraph("Téléphone : {{tenant_phone}}")
    doc.add_paragraph("Courriel : {{tenant_email}}")
    doc.add_paragraph()
    
    doc.add_paragraph("SOUS-LOCATAIRE :")
    doc.add_paragraph("Nom : {{subtenant_name}}")
    doc.add_paragraph("Adresse actuelle : {{subtenant_current_address}}")
    doc.add_paragraph("Téléphone : {{subtenant_phone}}")
    doc.add_paragraph("Courriel : {{subtenant_email}}")
    doc.add_paragraph()
    
    # Premises
    doc.add_paragraph("1. LIEUX LOUÉS")
    doc.add_paragraph("Le Sous-locateur accepte de sous-louer au Sous-locataire, et le Sous-locataire accepte de louer du Sous-locateur, les lieux situés à {{property_address}}, {{property_city}}, QC {{property_postal_code}} (« les Lieux »), comprenant : {{premises_description}}.")
    doc.add_paragraph()
    
    # Term
    doc.add_paragraph("2. DURÉE")
    doc.add_paragraph("La durée de cette sous-location commencera le {{start_date}} et se terminera le {{end_date}} (« Durée »).")
    doc.add_paragraph()
    
    # Rent
    doc.add_paragraph("3. LOYER")
    doc.add_paragraph("Le Sous-locataire versera un loyer mensuel de {{monthly_rent}} $ CAD, payable le {{rent_due_day}} jour de chaque mois. Le premier paiement est dû le {{first_payment_date}}.")
    doc.add_paragraph()
    
    doc.add_paragraph("Mode de paiement : {{payment_method}}")
    doc.add_paragraph()
    
    # Security deposit
    doc.add_paragraph("4. DÉPÔT DE GARANTIE")
    doc.add_paragraph("Le Sous-locataire versera un dépôt de garantie de {{security_deposit}} $ CAD à la signature de ce contrat. Ce dépôt sera remboursé dans les 10 jours suivant la fin de la sous-location, moins toute déduction pour dommages dépassant l'usure normale ou loyers impayés.")
    doc.add_paragraph()
    
    # Utilities
    doc.add_paragraph("5. SERVICES PUBLICS")
    doc.add_paragraph("Les services publics suivants sont inclus dans le loyer : {{utilities_included}}.")
    doc.add_paragraph("Les services publics suivants sont la responsabilité du Sous-locataire : {{utilities_not_included}}.")
    doc.add_paragraph()
    
    # Use of premises
    doc.add_paragraph("6. UTILISATION DES LIEUX")
    doc.add_paragraph("Les Lieux doivent être utilisés uniquement comme résidence privée pour le(s) occupant(s) nommé(s) ci-dessus. Le Sous-locataire ne doit pas sous-louer les Lieux à une autre partie.")
    doc.add_paragraph()
    
    # Original lease
    doc.add_paragraph("7. BAIL PRINCIPAL")
    doc.add_paragraph("Cette sous-location est subordonnée au bail principal entre le Sous-locateur et {{landlord_name}} (« Propriétaire »), daté du {{original_lease_date}}. Une copie du bail principal est jointe à ce contrat, et le Sous-locataire accepte de se conformer à toutes les conditions du bail principal.")
    doc.add_paragraph()
    
    doc.add_paragraph("Conformément au Code civil du Québec, qui permet la sous-location avec notification au propriétaire, le Sous-locateur confirme avoir envoyé un avis de sous-location au Propriétaire le {{landlord_notice_date}}.")
    doc.add_paragraph()
    
    # Tenant obligations
    doc.add_paragraph("8. OBLIGATIONS DU SOUS-LOCATAIRE")
    doc.add_paragraph("Le Sous-locataire accepte de :")
    doc.add_paragraph("a) Maintenir les Lieux dans un état propre et sanitaire.")
    doc.add_paragraph("b) Se conformer à toutes les obligations qui s'appliqueraient au Sous-locateur selon le bail principal.")
    doc.add_paragraph("c) Respecter les règles de l'immeuble et les règlements de l'association des propriétaires (le cas échéant).")
    doc.add_paragraph("d) Permettre au Sous-locateur d'accéder aux Lieux avec un préavis raisonnable pour inspection, réparations ou pour montrer les Lieux à de futurs locataires.")
    doc.add_paragraph()
    
    # Lead-based paint disclosure
    doc.add_paragraph("9. AUTRES DISPOSITIONS")
    doc.add_paragraph("{{additional_provisions}}")
    doc.add_paragraph()
    
    # Governing law
    doc.add_paragraph("10. LOI APPLICABLE")
    doc.add_paragraph("Ce contrat est régi par les lois de la province de Québec et le Code civil du Québec.")
    doc.add_paragraph()
    
    # Signatures
    doc.add_paragraph("LES PARTIES ONT EXÉCUTÉ CE CONTRAT À LA DATE INDIQUÉE CI-DESSUS.")
    doc.add_paragraph()
    doc.add_paragraph("________________________")
    doc.add_paragraph("Sous-locateur: {{tenant_name}}")
    doc.add_paragraph("Date: ___________________")
    doc.add_paragraph()
    doc.add_paragraph("________________________")
    doc.add_paragraph("Sous-locataire: {{subtenant_name}}")
    doc.add_paragraph("Date: ___________________")
    doc.add_paragraph()
    
    # Landlord consent (if applicable)
    doc.add_paragraph("CONSENTEMENT DU PROPRIÉTAIRE (le cas échéant):")
    doc.add_paragraph()
    doc.add_paragraph("Je, {{landlord_name}}, propriétaire des Lieux, consent à cette sous-location.")
    doc.add_paragraph()
    doc.add_paragraph("________________________")
    doc.add_paragraph("Propriétaire: {{landlord_name}}")
    doc.add_paragraph("Date: ___________________")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_sublease_agreement_template(province_code=None):
    """Create a specialized template for sublease agreement"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    # Create a base template name
    base_template_path = f"{template_dir}/sublease_agreement.docx"
    
    # If a province code is provided, create a province-specific template
    if province_code:
        output_path = f"{template_dir}/sublease_agreement_{province_code}.docx"
        legislation = get_provincial_legislation(province_code)
        sublease_restrictions = get_sublease_restrictions(province_code)
    else:
        output_path = base_template_path
        legislation = "Residential Tenancies Act"
        sublease_restrictions = "may require written consent from the landlord depending on your specific lease and local regulations"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('SUBLEASE AGREEMENT', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    doc.add_paragraph()
    
    # Add introduction
    doc.add_paragraph('THIS SUBLEASE AGREEMENT ("Agreement") is made between:')
    doc.add_paragraph()
    
    # Add parties
    doc.add_paragraph('ORIGINAL TENANT ("Sublandlord"):')
    doc.add_paragraph("Name: {{tenant_name}}")
    doc.add_paragraph("Current Address: {{property_address}}, {{property_city}}, {{property_province}} {{property_postal_code}}")
    doc.add_paragraph("Phone: {{tenant_phone}}")
    doc.add_paragraph("Email: {{tenant_email}}")
    doc.add_paragraph()
    
    doc.add_paragraph("SUBTENANT:")
    doc.add_paragraph("Name: {{subtenant_name}}")
    doc.add_paragraph("Current Address: {{subtenant_current_address}}")
    doc.add_paragraph("Phone: {{subtenant_phone}}")
    doc.add_paragraph("Email: {{subtenant_email}}")
    doc.add_paragraph()
    
    # Premises
    doc.add_paragraph("1. PREMISES")
    doc.add_paragraph('Sublandlord agrees to sublease to Subtenant, and Subtenant agrees to sublease from Sublandlord, the premises located at {{property_address}}, {{property_city}}, {{property_province}} {{property_postal_code}} (the "Premises"), consisting of: {{premises_description}}.')
    doc.add_paragraph()
    
    # Term
    doc.add_paragraph("2. TERM")
    doc.add_paragraph('The term of this sublease shall begin on {{start_date}} and end on {{end_date}} (the "Term").')
    doc.add_paragraph()
    
    # Rent
    doc.add_paragraph("3. RENT")
    doc.add_paragraph("Subtenant shall pay a monthly rent of ${{monthly_rent}} CAD, due on the {{rent_due_day}} day of each month. The first payment is due on {{first_payment_date}}.")
    doc.add_paragraph()
    
    doc.add_paragraph("Payment Method: {{payment_method}}")
    doc.add_paragraph()
    
    # Security deposit
    doc.add_paragraph("4. SECURITY DEPOSIT")
    doc.add_paragraph("Subtenant shall pay a security deposit of ${{security_deposit}} CAD upon signing this Agreement. This deposit will be returned within 10 days after the end of the sublease term, less any deductions for damages beyond normal wear and tear or unpaid rent.")
    doc.add_paragraph()
    
    # Utilities
    doc.add_paragraph("5. UTILITIES")
    doc.add_paragraph("The following utilities are included in the rent: {{utilities_included}}.")
    doc.add_paragraph("The following utilities are the responsibility of the Subtenant: {{utilities_not_included}}.")
    doc.add_paragraph()
    
    # Use of premises
    doc.add_paragraph("6. USE OF PREMISES")
    doc.add_paragraph("The Premises shall be used for residential purposes only by the Subtenant(s) named above. Subtenant shall not further sublet the Premises to another party.")
    doc.add_paragraph()
    
    # Original lease
    doc.add_paragraph("7. MASTER LEASE")
    doc.add_paragraph(f'This sublease is subject to the terms and conditions of the master lease between Sublandlord and {{landlord_name}} ("Landlord"), dated {{original_lease_date}}. A copy of the master lease is attached to this Agreement, and Subtenant agrees to comply with all terms of the master lease.')
    doc.add_paragraph()
    
    doc.add_paragraph(f"Under the {legislation}, subleasing {sublease_restrictions}. Sublandlord confirms that: ")
    if province_code == "QC":
        doc.add_paragraph("☐ Landlord has been notified of this sublease on {{landlord_notice_date}}")
    else:
        doc.add_paragraph("☐ Landlord has provided written consent for this sublease on {{landlord_consent_date}}")
        doc.add_paragraph("☐ Landlord's consent is not required based on the terms of the master lease")
    doc.add_paragraph()
    
    # Tenant obligations
    doc.add_paragraph("8. SUBTENANT OBLIGATIONS")
    doc.add_paragraph("Subtenant agrees to:")
    doc.add_paragraph("a) Keep the Premises in a clean and sanitary condition.")
    doc.add_paragraph("b) Comply with all obligations that would apply to the Sublandlord under the master lease.")
    doc.add_paragraph("c) Respect building rules and homeowner association regulations (if applicable).")
    doc.add_paragraph("d) Allow Sublandlord access to the Premises with reasonable notice for inspection, repairs, or to show the Premises to prospective tenants.")
    doc.add_paragraph()
    
    # Additional provisions
    doc.add_paragraph("9. ADDITIONAL PROVISIONS")
    doc.add_paragraph("{{additional_provisions}}")
    doc.add_paragraph()
    
    # Governing law
    doc.add_paragraph("10. GOVERNING LAW")
    doc.add_paragraph(f"This Agreement is governed by the laws of the province of {{property_province}} and the {legislation}.")
    doc.add_paragraph()
    
    # Signatures
    doc.add_paragraph("THE PARTIES HAVE EXECUTED THIS AGREEMENT AS OF THE DATE FIRST WRITTEN ABOVE.")
    doc.add_paragraph()
    doc.add_paragraph("________________________")
    doc.add_paragraph("Sublandlord: {{tenant_name}}")
    doc.add_paragraph("Date: ___________________")
    doc.add_paragraph()
    doc.add_paragraph("________________________")
    doc.add_paragraph("Subtenant: {{subtenant_name}}")
    doc.add_paragraph("Date: ___________________")
    doc.add_paragraph()
    
    # Landlord consent (if applicable)
    doc.add_paragraph("LANDLORD CONSENT (if applicable):")
    doc.add_paragraph()
    doc.add_paragraph("I, {{landlord_name}}, landlord of the Premises, consent to this sublease.")
    doc.add_paragraph()
    doc.add_paragraph("________________________")
    doc.add_paragraph("Landlord: {{landlord_name}}")
    doc.add_paragraph("Date: ___________________")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_all_provincial_templates():
    """Create sublease agreement templates for all provinces"""
    provinces = ["ON", "BC", "AB", "QC", "MB", "SK", "NS", "NB", "NL", "PE", "NT", "NU", "YT"]
    template_count = 0
    
    # First create generic template
    create_sublease_agreement_template()
    
    # Then create province-specific templates
    for province in provinces:
        create_sublease_agreement_template(province)
        template_count += 1
    
    # Create Quebec French template
    create_french_quebec_template()
    
    print(f"Created {template_count} provincial sublease agreement templates.")

if __name__ == "__main__":
    create_all_provincial_templates()