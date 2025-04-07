#!/usr/bin/env python3
"""
Template Creation Script for Tenant Notice Documents
Creates templates for:
1. Repair Notice
2. Intent to Vacate Notice 
3. Termination Notice
"""

import os
import sys
from docxtpl import DocxTemplate
from docx import Document
import docx

def get_provincial_legislation(province_code):
    """Return the appropriate legislation name based on province code"""
    legislation_map = {
        "ON": "Residential Tenancies Act",
        "BC": "Residential Tenancy Act",
        "AB": "Residential Tenancy Act",
        "QC": "Civil Code of Quebec",
        "MB": "Residential Tenancies Act",
        "SK": "Residential Tenancies Act",
        "NS": "Residential Tenancies Act",
        "NB": "Residential Tenancies Act",
        "NL": "Residential Tenancies Act",
        "PE": "Rental of a Residential Property Act",
        "NT": "Residential Tenancies Act",
        "NU": "Residential Tenancies Act",
        "YT": "Residential Landlord and Tenant Act"
    }
    
    return legislation_map.get(province_code, "Residential Tenancies Act")

def get_repair_timeline(province_code):
    """Return province-specific reasonable repair timeline guidance"""
    timeline_map = {
        "ON": "14 days for regular repairs, 24 hours for emergency repairs",
        "BC": "30 days for regular repairs, 24 hours for emergency repairs",
        "AB": "14 days for regular repairs, immediately for emergency repairs",
        "QC": "10 days for regular repairs, 24 hours for emergency repairs",
        "MB": "7 days for regular repairs, 24 hours for emergency repairs",
        "SK": "14 days for regular repairs, immediately for emergency repairs",
        "NS": "7 days for regular repairs, 24 hours for emergency repairs",
        "NB": "7 days for regular repairs, 24 hours for emergency repairs",
        "NL": "14 days for regular repairs, 24 hours for emergency repairs",
        "PE": "7 days for regular repairs, 24 hours for emergency repairs",
        "NT": "7 days for regular repairs, 24 hours for emergency repairs", 
        "NU": "7 days for regular repairs, 24 hours for emergency repairs",
        "YT": "7 days for regular repairs, 24 hours for emergency repairs"
    }
    
    return timeline_map.get(province_code, "14 days for regular repairs, 24 hours for emergency repairs")

def get_notice_period(province_code, notice_type="vacate"):
    """Return province-specific notice periods for vacating or termination"""
    # Notice periods for standard end-of-lease vacating
    vacate_map = {
        "ON": "60 days for month-to-month tenancies",
        "BC": "30 days for month-to-month tenancies",
        "AB": "30 days for month-to-month tenancies",
        "QC": "30 days for leases less than 6 months, 60 days for 6+ months",
        "MB": "30 days for month-to-month tenancies",
        "SK": "30 days for month-to-month tenancies",
        "NS": "30 days for month-to-month tenancies",
        "NB": "30 days for month-to-month tenancies",
        "NL": "30 days for month-to-month tenancies",
        "PE": "60 days for month-to-month tenancies",
        "NT": "30 days for month-to-month tenancies",
        "NU": "30 days for month-to-month tenancies",
        "YT": "30 days for month-to-month tenancies"
    }
    
    # Notice periods for early termination (exceptional circumstances)
    termination_map = {
        "ON": "28 days for domestic violence, 60 days for health reasons",
        "BC": "30 days for domestic violence or landlord breach",
        "AB": "28 days for domestic violence, immediate for safety threats",
        "QC": "30 days for senior home relocation, 60 days for health reasons",
        "MB": "14 days for domestic violence, 30 days for health reasons",
        "SK": "28 days for domestic violence, immediate for safety threats",
        "NS": "30 days for health reasons or domestic violence",
        "NB": "30 days for domestic violence, immediate for health/safety",
        "NL": "30 days for domestic violence, immediate for safety threats",
        "PE": "60 days for health reasons, 30 days for domestic violence",
        "NT": "30 days for health reasons or domestic violence",
        "NU": "30 days for health reasons or domestic violence",
        "YT": "28 days for domestic violence, immediate for safety threats"
    }
    
    if notice_type == "vacate":
        return vacate_map.get(province_code, "30 days for month-to-month tenancies")
    else:  # termination
        return termination_map.get(province_code, "30 days for exceptional circumstances")

def create_repair_notice_template(province_code=None):
    """Create a specialized template for repair notice letters"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    # Create a base template name
    base_template_path = f"{template_dir}/repair_notice.docx"
    
    # If a province code is provided, create a province-specific template
    if province_code:
        output_path = f"{template_dir}/repair_notice_{province_code}.docx"
        legislation = get_provincial_legislation(province_code)
        repair_timeline = get_repair_timeline(province_code)
    else:
        output_path = base_template_path
        legislation = "Residential Tenancies Act"
        repair_timeline = "14 days for regular repairs, 24 hours for emergency repairs"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('NOTICE OF REPAIR REQUEST', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add landlord info
    doc.add_paragraph("To:")
    doc.add_paragraph("{{landlord_name}}")
    doc.add_paragraph("{{landlord_address}}")
    doc.add_paragraph("{{landlord_city}}, {{landlord_province}} {{landlord_postal_code}}")
    doc.add_paragraph()
    
    # Add tenant info
    doc.add_paragraph("From:")
    doc.add_paragraph("{{tenant_name}}")
    doc.add_paragraph("{{property_address}}")
    doc.add_paragraph("{{property_city}}, {{property_province}} {{property_postal_code}}")
    doc.add_paragraph()
    
    # Add rental property reference
    doc.add_paragraph("Re: Repair Request for Rental Property at {{property_address}}, {{property_city}}, {{property_province}} {{property_postal_code}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph(f"Dear {{landlord_name}},")
    doc.add_paragraph()
    
    doc.add_paragraph("This letter serves as a formal notice to request repairs to the rental property identified above, as required under the " + legislation + ".")
    doc.add_paragraph()
    
    doc.add_paragraph("The following repairs are needed:")
    doc.add_paragraph("{{repair_description}}")
    doc.add_paragraph()
    
    # Add specific details about safety or habitability issues if they exist
    doc.add_paragraph("Impact on the property and tenants:")
    doc.add_paragraph("{{impact_description}}")
    doc.add_paragraph()
    
    # Add reference to previous communications if any
    doc.add_paragraph("Reference to previous communications (if applicable):")
    doc.add_paragraph("{{previous_communications}}")
    doc.add_paragraph()
    
    # Add timeline for repairs
    doc.add_paragraph(f"Under the {legislation}, landlords are generally expected to address necessary repairs within a reasonable timeframe ({repair_timeline}).")
    doc.add_paragraph()
    
    # Add deadline request
    doc.add_paragraph("I respectfully request that these repairs be completed by {{requested_completion_date}}.")
    doc.add_paragraph()
    
    # Add inspection access information
    doc.add_paragraph("I am available for you or your maintenance staff to access the property on the following dates and times:")
    doc.add_paragraph("{{available_dates}}")
    doc.add_paragraph()
    
    # Add closing paragraph with legal rights reservation
    doc.add_paragraph(f"If these repairs are not addressed within the specified timeframe, I reserve the right to pursue remedies available under the {legislation}, which may include filing a complaint with the local tenancy authority, withholding rent (where legally permitted), or terminating the lease agreement.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{tenant_name}}")
    doc.add_paragraph("Phone: {{tenant_phone}}")
    doc.add_paragraph("Email: {{tenant_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_intent_to_vacate_template(province_code=None):
    """Create a specialized template for intent to vacate letter"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    # Create a base template name
    base_template_path = f"{template_dir}/intent_to_vacate.docx"
    
    # If a province code is provided, create a province-specific template
    if province_code:
        output_path = f"{template_dir}/intent_to_vacate_{province_code}.docx"
        legislation = get_provincial_legislation(province_code)
        notice_period = get_notice_period(province_code, "vacate")
    else:
        output_path = base_template_path
        legislation = "Residential Tenancies Act"
        notice_period = "30 days for month-to-month tenancies"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('NOTICE OF INTENT TO VACATE', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add landlord info
    doc.add_paragraph("To:")
    doc.add_paragraph("{{landlord_name}}")
    doc.add_paragraph("{{landlord_address}}")
    doc.add_paragraph("{{landlord_city}}, {{landlord_province}} {{landlord_postal_code}}")
    doc.add_paragraph()
    
    # Add tenant info
    doc.add_paragraph("From:")
    doc.add_paragraph("{{tenant_name}}")
    doc.add_paragraph("{{property_address}}")
    doc.add_paragraph("{{property_city}}, {{property_province}} {{property_postal_code}}")
    doc.add_paragraph()
    
    # Add rental property reference
    doc.add_paragraph("Re: Notice to Vacate Rental Property at {{property_address}}, {{property_city}}, {{property_province}} {{property_postal_code}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph(f"Dear {{landlord_name}},")
    doc.add_paragraph()
    
    doc.add_paragraph(f"This letter serves as my formal notice of intent to vacate the rental property identified above, in accordance with the {legislation} which requires {notice_period}.")
    doc.add_paragraph()
    
    doc.add_paragraph("I will be vacating the premises on {{vacate_date}}, which provides the required notice period as stipulated in our lease agreement and provincial law.")
    doc.add_paragraph()
    
    # Add request for move-out inspection
    doc.add_paragraph("I request that we schedule a move-out inspection prior to my departure. I am available on the following dates and times:")
    doc.add_paragraph("{{inspection_dates}}")
    doc.add_paragraph()
    
    # Add forwarding information
    doc.add_paragraph("Please forward my security deposit to the following address:")
    doc.add_paragraph("{{forwarding_address}}")
    doc.add_paragraph("{{forwarding_city}}, {{forwarding_province}} {{forwarding_postal_code}}")
    doc.add_paragraph()
    
    # Add utility information if applicable
    doc.add_paragraph("I will arrange for the disconnection or transfer of utilities on {{utility_date}}.")
    doc.add_paragraph()
    
    # Add closing paragraph
    doc.add_paragraph("I have appreciated my time as a tenant at this property. Please contact me at your earliest convenience to confirm receipt of this notice and to arrange the move-out inspection.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{tenant_name}}")
    doc.add_paragraph("Phone: {{tenant_phone}}")
    doc.add_paragraph("Email: {{tenant_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_termination_notice_template(province_code=None):
    """Create a specialized template for lease termination letter"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    # Create a base template name
    base_template_path = f"{template_dir}/termination_notice.docx"
    
    # If a province code is provided, create a province-specific template
    if province_code:
        output_path = f"{template_dir}/termination_notice_{province_code}.docx"
        legislation = get_provincial_legislation(province_code)
        notice_period = get_notice_period(province_code, "termination")
    else:
        output_path = base_template_path
        legislation = "Residential Tenancies Act"
        notice_period = "30 days for exceptional circumstances"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('NOTICE OF LEASE TERMINATION', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add landlord info
    doc.add_paragraph("To:")
    doc.add_paragraph("{{landlord_name}}")
    doc.add_paragraph("{{landlord_address}}")
    doc.add_paragraph("{{landlord_city}}, {{landlord_province}} {{landlord_postal_code}}")
    doc.add_paragraph()
    
    # Add tenant info
    doc.add_paragraph("From:")
    doc.add_paragraph("{{tenant_name}}")
    doc.add_paragraph("{{property_address}}")
    doc.add_paragraph("{{property_city}}, {{property_province}} {{property_postal_code}}")
    doc.add_paragraph()
    
    # Add rental property reference
    doc.add_paragraph("Re: Immediate Termination of Lease for Rental Property at {{property_address}}, {{property_city}}, {{property_province}} {{property_postal_code}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph(f"Dear {{landlord_name}},")
    doc.add_paragraph()
    
    doc.add_paragraph(f"This letter serves as my formal notice of lease termination for the rental property identified above, in accordance with the {legislation} which allows for early termination under exceptional circumstances ({notice_period}).")
    doc.add_paragraph()
    
    doc.add_paragraph("I am terminating my lease effective {{termination_date}} due to the following circumstances:")
    doc.add_paragraph("{{termination_reason}}")
    doc.add_paragraph()
    
    # Add supporting documentation reference
    doc.add_paragraph("Supporting documentation (if applicable):")
    doc.add_paragraph("{{supporting_documentation}}")
    doc.add_paragraph()
    
    # Add request for move-out inspection
    doc.add_paragraph("I request that we schedule a move-out inspection prior to my departure. I am available on the following dates and times:")
    doc.add_paragraph("{{inspection_dates}}")
    doc.add_paragraph()
    
    # Add forwarding information
    doc.add_paragraph("Please forward my security deposit to the following address:")
    doc.add_paragraph("{{forwarding_address}}")
    doc.add_paragraph("{{forwarding_city}}, {{forwarding_province}} {{forwarding_postal_code}}")
    doc.add_paragraph()
    
    # Add legal reference for special circumstances termination
    doc.add_paragraph(f"As per the {legislation}, my circumstances meet the criteria for early lease termination. I have provided the appropriate notice as required by law for these special circumstances.")
    doc.add_paragraph()
    
    # Add closing paragraph
    doc.add_paragraph("Please contact me at your earliest convenience to confirm receipt of this notice and to arrange the move-out inspection.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{tenant_name}}")
    doc.add_paragraph("Phone: {{tenant_phone}}")
    doc.add_paragraph("Email: {{tenant_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_all_provincial_templates():
    """Create templates for all provinces"""
    provinces = ["ON", "BC", "AB", "QC", "MB", "SK", "NS", "NB", "NL", "PE", "NT", "NU", "YT"]
    repair_count = 0
    vacate_count = 0
    termination_count = 0
    
    # First create generic templates
    create_repair_notice_template()
    create_intent_to_vacate_template()
    create_termination_notice_template()
    
    # Then create province-specific templates
    for province in provinces:
        create_repair_notice_template(province)
        repair_count += 1
        
        create_intent_to_vacate_template(province)
        vacate_count += 1
        
        create_termination_notice_template(province)
        termination_count += 1
    
    print(f"Created {repair_count} provincial repair notice templates.")
    print(f"Created {vacate_count} provincial intent to vacate templates.")
    print(f"Created {termination_count} provincial termination notice templates.")

if __name__ == "__main__":
    create_all_provincial_templates()