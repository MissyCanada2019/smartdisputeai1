from flask import Flask, request, render_template, send_from_directory, redirect, url_for
from docxtpl import DocxTemplate
from docx import Document
import os
import uuid
import json
import datetime
from flask import jsonify

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'smartdispute_ai_development_key')

def get_legal_act(dispute_type, province):
    """Get the appropriate legal act based on dispute type and province"""
    legal_acts = {
        'landlord_tenant': {
            'ON': 'Residential Tenancies Act',
            'BC': 'Residential Tenancy Act',
            'AB': 'Residential Tenancies Act',
            'QC': 'Civil Code of Québec (sections governing residential leases)',
            'MB': 'Residential Tenancies Act',
            'SK': 'Residential Tenancies Act',
            'NS': 'Residential Tenancies Act',
            'NB': 'Residential Tenancies Act',
            'NL': 'Residential Tenancies Act',
            'PE': 'Rental of Residential Property Act',
            'NT': 'Residential Tenancies Act',
            'NU': 'Residential Tenancies Act',
            'YT': 'Residential Landlord and Tenant Act'
        },
        'landlord_cease_desist': {
            'ON': 'Residential Tenancies Act',
            'BC': 'Residential Tenancy Act',
            'AB': 'Residential Tenancies Act',
            'QC': 'Civil Code of Québec (sections governing residential leases)',
            'MB': 'Residential Tenancies Act',
            'SK': 'Residential Tenancies Act',
            'NS': 'Residential Tenancies Act',
            'NB': 'Residential Tenancies Act',
            'NL': 'Residential Tenancies Act',
            'PE': 'Rental of Residential Property Act',
            'NT': 'Residential Tenancies Act',
            'NU': 'Residential Tenancies Act',
            'YT': 'Residential Landlord and Tenant Act'
        },
        'credit': {
            'ON': 'Consumer Reporting Act',
            'BC': 'Business Practices and Consumer Protection Act',
            'AB': 'Consumer Protection Act',
            'QC': 'Act Respecting the Protection of Personal Information in the Private Sector',
            'MB': 'Consumer Protection Act',
            'SK': 'Consumer Protection and Business Practices Act',
            'NS': 'Consumer Reporting Act',
            'NB': 'Consumer Reporting Act',
            'NL': 'Consumer Protection and Business Practices Act',
            'PE': 'Consumer Reporting Act',
            'NT': 'Consumer Protection Act',
            'NU': 'Consumer Protection Act',
            'YT': 'Consumer Protection Act'
        },
        'cas': {
            'ON': 'Child, Youth and Family Services Act',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act (Loi sur la protection de la jeunesse)',
            'MB': 'Child and Family Services Act',
            'SK': 'Child and Family Services Act',
            'NS': 'Children and Family Services Act',
            'NB': 'Family Services Act',
            'NL': 'Children, Youth and Families Act',
            'PE': 'Child Protection Act',
            'NT': 'Child and Family Services Act',
            'NU': 'Child and Family Services Act',
            'YT': 'Child and Family Services Act'
        },
        'cas_cease_desist': {
            'ON': 'Child, Youth and Family Services Act',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act (Loi sur la protection de la jeunesse)',
            'MB': 'Child and Family Services Act',
            'SK': 'Child and Family Services Act',
            'NS': 'Children and Family Services Act',
            'NB': 'Family Services Act',
            'NL': 'Children, Youth and Families Act',
            'PE': 'Child Protection Act',
            'NT': 'Child and Family Services Act',
            'NU': 'Child and Family Services Act',
            'YT': 'Child and Family Services Act'
        },
        'cas_worker_reassignment': {
            'ON': 'Child, Youth and Family Services Act, 2017',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act',
            'MB': 'Child and Family Services Act (Manitoba)',
            'SK': 'The Child and Family Services Act',
            'NS': 'Children and Family Services Act (Nova Scotia)',
            'NB': 'Family Services Act (New Brunswick)',
            'NL': 'Children, Youth and Families Act (Newfoundland and Labrador)',
            'PE': 'Child Protection Act (Prince Edward Island)',
            'NT': 'Child and Family Services Act (Northwest Territories)',
            'YT': 'Child and Family Services Act (Yukon)',
            'NU': 'Child and Family Services Act (Nunavut)'
        }
    }
    
    return legal_acts.get(dispute_type, {}).get(province, 'applicable provincial legislation')

def get_agency_name(province):
    """Get the appropriate child services agency name based on province code"""
    agency_map = {
        "ON": "Children's Aid Society",
        "BC": "Ministry of Children and Family Development",
        "AB": "Children's Services",
        "QC": "Direction de la protection de la jeunesse (DPJ)",
        "MB": "Child and Family Services",
        "SK": "Ministry of Social Services",
        "NS": "Department of Community Services",
        "NB": "Department of Social Development",
        "NL": "Department of Children, Seniors and Social Development",
        "PE": "Child Protection Services",
        "NT": "Child and Family Services",
        "NU": "Department of Family Services",
        "YT": "Family and Children's Services"
    }
    
    return agency_map.get(province, "Child Protection Agency")
UPLOAD_FOLDER = 'generated'
TEMPLATE_FOLDER = 'templates/docs'

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

if not os.path.exists(TEMPLATE_FOLDER):
    os.makedirs(TEMPLATE_FOLDER)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/form')
def dynamic_form():
    province = request.args.get('province')
    dispute_type = request.args.get('type')
    return render_template('form.html', province=province, dispute_type=dispute_type)

@app.route('/generate', methods=['POST'])
def generate():
    name = request.form['name']
    email = request.form['email']
    description = request.form['description']
    province = request.form['province']
    dispute_type = request.form['dispute_type']
    
    # Get optional form fields with defaults
    recipient_name = request.form.get('recipient_name', 'To Whom It May Concern')
    recipient_address = request.form.get('recipient_address', '')
    user_address = request.form.get('address', '')
    
    # Language preference for Quebec (default to English)
    language = request.form.get('language', 'EN')
    
    # Get template filename based on dispute type and province
    filename = f"{dispute_type}_{province}.docx"
    
    # Special handling for Quebec French templates
    if province == 'QC' and language == 'FR' and dispute_type == 'cas_cease_desist':
        filename = f"{dispute_type}_{province}_FR.docx"
    
    template_path = os.path.join(TEMPLATE_FOLDER, filename)

    # Select appropriate template creation function based on dispute type
    if not os.path.exists(template_path):
        if dispute_type == 'landlord_cease_desist':
            try:
                # Try to import and use the specialized template creator
                from create_landlord_cease_desist_template import create_landlord_cease_desist_template
                create_landlord_cease_desist_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating landlord cease & desist template: {e}")
                create_basic_template(template_path)
        elif dispute_type == 'cas_cease_desist':
            try:
                # Try to import and use the specialized template creator
                from create_cas_cease_desist_template import create_cas_cease_desist_template, create_french_quebec_template
                if province == 'QC' and language == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_cease_desist_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS cease & desist template: {e}")
                create_basic_template(template_path)
        elif dispute_type == 'cas_worker_reassignment':
            try:
                # Try to import and use the specialized template creator
                from create_cas_worker_reassignment_template import create_cas_worker_reassignment_template, create_french_quebec_template
                if province == 'QC' and language == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_worker_reassignment_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS worker reassignment template: {e}")
                create_basic_template(template_path)
        else:
            # Fall back to basic template
            create_basic_template(template_path)

    # Get legal act based on dispute type and province
    legal_act = get_legal_act(dispute_type, province)
    
    # Get agency name for CAS templates
    agency_name = get_agency_name(province) if dispute_type.startswith('cas') else ''

    # Get cease and desist specific fields if applicable
    prior_incidents = request.form.get('prior_incidents', '')
    requested_actions = request.form.get('requested_actions', '')
    consequences = request.form.get('consequences', '')
    
    # Get worker reassignment specific fields if applicable
    current_worker_name = request.form.get('current_worker_name', '')
    description_of_concerns = request.form.get('description_of_concerns', '')
    previous_attempts = request.form.get('previous_attempts', '')
    payment_method = request.form.get('payment_method', '')
    
    # Check if payment is required
    requires_payment = dispute_type == 'cas_worker_reassignment'
    
    # If payment is required but payment processing is not completed
    if requires_payment and not request.args.get('payment_completed'):
        # Store form data in session
        session_data = {
            'name': name,
            'email': email,
            'description': description,
            'province': province,
            'dispute_type': dispute_type,
            'recipient_name': recipient_name,
            'recipient_address': recipient_address,
            'user_address': user_address,
            'current_worker_name': current_worker_name,
            'description_of_concerns': description_of_concerns,
            'previous_attempts': previous_attempts,
            'payment_method': payment_method
        }
        
        # Save session data to a temporary file
        session_id = str(uuid.uuid4())
        session_path = os.path.join(UPLOAD_FOLDER, f"session_{session_id}.json")
        with open(session_path, 'w') as f:
            json.dump(session_data, f)
        
        # Redirect to payment page
        return redirect(url_for('payment', 
                               session_id=session_id, 
                               amount='9.99', 
                               payment_method=payment_method))
    
    doc = DocxTemplate(template_path)
    context = {
        'name': name,
        'email': email,
        'description': description,
        'province': province,
        'recipient_name': recipient_name,
        'recipient_address': recipient_address,
        'address': user_address,
        'legal_act': legal_act,
        'agency_name': agency_name,
        'prior_incidents': prior_incidents,
        'requested_actions': requested_actions, 
        'consequences': consequences,
        'current_worker_name': current_worker_name,
        'description_of_concerns': description_of_concerns,
        'previous_attempts': previous_attempts,
        'now': lambda: datetime.datetime.now().strftime("%B %d, %Y")
    }

    doc.render(context)
    output_filename = f"{uuid.uuid4()}.docx"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)
    doc.save(output_path)

    return redirect(url_for('success', filename=output_filename))

def create_basic_template(path):
    """Create a basic template if none exists"""
    doc = Document()
    doc.add_heading('SmartDispute.ai Generated Document', 0)
    
    # Sender information
    doc.add_heading('Sender Information', level=1)
    doc.add_paragraph('Name: {{ name }}')
    doc.add_paragraph('Email: {{ email }}')
    if '{{ address }}':
        doc.add_paragraph('Address: {{ address }}')
    doc.add_paragraph('Date: {{ now() }}')
    
    # Recipient information
    doc.add_heading('Recipient Information', level=1)
    doc.add_paragraph('To: {{ recipient_name }}')
    if '{{ recipient_address }}':
        doc.add_paragraph('{{ recipient_address }}')
    
    # Legal information
    doc.add_heading('Legal References', level=1)
    doc.add_paragraph('Province: {{ province }}')
    if '{{ legal_act }}':
        doc.add_paragraph('Applicable Legislation: {{ legal_act }}')
    if '{{ agency_name }}':
        doc.add_paragraph('Agency: {{ agency_name }}')
    
    # Issue description
    doc.add_heading('Issue Description', level=1)
    doc.add_paragraph('{{ description }}')
    
    # Cease and desist specific sections
    if path.find('cease_desist') != -1:
        doc.add_heading('Prior Incidents', level=1)
        doc.add_paragraph('{{ prior_incidents }}')
        
        doc.add_heading('Requested Actions', level=1)
        doc.add_paragraph('{{ requested_actions }}')
        
        doc.add_heading('Consequences', level=1)
        doc.add_paragraph('{{ consequences }}')
    
    # Worker reassignment specific sections
    if path.find('worker_reassignment') != -1:
        doc.add_heading('Current CAS Worker', level=1)
        doc.add_paragraph('{{ current_worker_name }}')
        
        doc.add_heading('Concerns', level=1)
        doc.add_paragraph('{{ description_of_concerns }}')
        
        doc.add_heading('Previous Attempts to Resolve', level=1)
        doc.add_paragraph('{{ previous_attempts }}')
    
    # Signature
    doc.add_paragraph('\n\nSincerely,\n\n\n\n{{ name }}')
    
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(path), exist_ok=True)
    
    doc.save(path)
    print(f"Created basic template at {path}")

@app.route('/payment')
def payment():
    """Handle payment page for paid documents"""
    session_id = request.args.get('session_id')
    amount = request.args.get('amount', '9.99')
    payment_method = request.args.get('payment_method', '')
    
    return render_template('payment.html', 
                           session_id=session_id, 
                           amount=amount, 
                           payment_method=payment_method)

@app.route('/process-payment', methods=['POST'])
def process_payment():
    """Process the payment and generate the document"""
    session_id = request.form.get('session_id')
    amount = request.form.get('amount')
    payment_method = request.form.get('payment_method')
    
    # In a real implementation, we would integrate with PayPal or a credit card processor here
    # For now, we'll simulate a successful payment
    
    # Load the stored session data
    session_path = os.path.join(UPLOAD_FOLDER, f"session_{session_id}.json")
    if not os.path.exists(session_path):
        return "Session expired or invalid. Please try again.", 400
    
    with open(session_path, 'r') as f:
        session_data = json.load(f)
    
    # Get the necessary data from the session
    name = session_data.get('name')
    email = session_data.get('email')
    description = session_data.get('description')
    province = session_data.get('province')
    dispute_type = session_data.get('dispute_type')
    recipient_name = session_data.get('recipient_name')
    recipient_address = session_data.get('recipient_address')
    user_address = session_data.get('user_address')
    current_worker_name = session_data.get('current_worker_name')
    description_of_concerns = session_data.get('description_of_concerns')
    previous_attempts = session_data.get('previous_attempts', '')
    
    # Get template filename and path
    filename = f"{dispute_type}_{province}.docx"
    template_path = os.path.join(TEMPLATE_FOLDER, filename)
    
    # Handle Quebec French template
    if province == 'QC' and session_data.get('language') == 'FR':
        filename = f"{dispute_type}_{province}_FR.docx"
        template_path = os.path.join(TEMPLATE_FOLDER, filename)
    
    # Make sure the template exists
    if not os.path.exists(template_path):
        try:
            # Try to import and use the specialized template creator
            from create_cas_worker_reassignment_template import create_cas_worker_reassignment_template, create_french_quebec_template
            
            if province == 'QC' and session_data.get('language') == 'FR':
                create_french_quebec_template()
            else:
                create_cas_worker_reassignment_template(province)
        except (ImportError, Exception) as e:
            print(f"Error creating CAS worker reassignment template: {e}")
            create_basic_template(template_path)
    
    # Get legal act and agency name
    legal_act = get_legal_act(dispute_type, province)
    agency_name = get_agency_name(province)
    
    # Create the document
    doc = DocxTemplate(template_path)
    context = {
        'name': name,
        'email': email,
        'description': description,
        'province': province,
        'recipient_name': recipient_name,
        'recipient_address': recipient_address,
        'address': user_address,
        'legal_act': legal_act,
        'agency_name': agency_name,
        'current_worker_name': current_worker_name,
        'description_of_concerns': description_of_concerns,
        'previous_attempts': previous_attempts,
        'now': lambda: datetime.datetime.now().strftime("%B %d, %Y")
    }
    
    # Render and save the document
    doc.render(context)
    output_filename = f"{uuid.uuid4()}.docx"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)
    doc.save(output_path)
    
    # Clean up the session file
    try:
        os.remove(session_path)
    except:
        pass
    
    # Redirect to success page
    return redirect(url_for('success', filename=output_filename, payment='completed'))

@app.route('/success')
def success():
    filename = request.args.get('filename')
    payment_completed = request.args.get('payment')
    return render_template('success.html', filename=filename, payment_completed=payment_completed)

@app.route('/generated/<filename>')
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5050)