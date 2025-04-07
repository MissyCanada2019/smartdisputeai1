"""
SmartDispute.ai - Dispute Letter Generator
Flask application for generating legal dispute letters based on templates
"""

import os
import re
import json
import uuid
import datetime
from flask import Flask, request, render_template, redirect, url_for, send_from_directory, jsonify
from docxtpl import DocxTemplate
from docx import Document
import pypandoc

app = Flask(__name__, template_folder='templates/flask')

# Configuration
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates', 'disputes')
GENERATED_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'generated')
STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static')

# Ensure directories exist
os.makedirs(TEMPLATE_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)

# Load legal reference data
LEGISLATION_MAP = {
    "ON": {
        "landlord_tenant": "Residential Tenancies Act, 2006",
        "credit_dispute": "Consumer Reporting Act (ON)",
        "cas": "Child, Youth and Family Services Act, 2017"
    },
    "BC": {
        "landlord_tenant": "Residential Tenancy Act",
        "credit_dispute": "Business Practices and Consumer Protection Act",
        "cas": "Child, Family and Community Service Act"
    },
    "AB": {
        "landlord_tenant": "Residential Tenancies Act",
        "credit_dispute": "Consumer Protection Act (AB)",
        "cas": "Child, Youth and Family Enhancement Act"
    },
    "QC": {
        "landlord_tenant": "Civil Code of Québec (articles 1851-1978)",
        "credit_dispute": "Consumer Protection Act (QC)",
        "cas": "Youth Protection Act"
    }
}

AUTHORITY_MAP = {
    "ON": {
        "landlord_tenant": "Landlord and Tenant Board",
        "credit_dispute": "Financial Services Regulatory Authority of Ontario",
        "cas": "Ontario Association of Children's Aid Societies"
    },
    "BC": {
        "landlord_tenant": "Residential Tenancy Branch",
        "credit_dispute": "Consumer Protection BC",
        "cas": "Ministry of Children and Family Development"
    },
    "AB": {
        "landlord_tenant": "Residential Tenancy Dispute Resolution Service",
        "credit_dispute": "Service Alberta, Consumer Investigations Unit",
        "cas": "Alberta Children's Services"
    },
    "QC": {
        "landlord_tenant": "Tribunal administratif du logement",
        "credit_dispute": "Office de la protection du consommateur",
        "cas": "Direction de la protection de la jeunesse"
    }
}

# Helper functions
def get_template_path(province, dispute_type):
    """Get path to correct template based on province and dispute type"""
    template_filename = f"template.docx"  # Default template name
    
    # Search for the template in order of specificity
    possible_paths = [
        os.path.join(TEMPLATE_DIR, province, dispute_type, template_filename),
        os.path.join(TEMPLATE_DIR, province, template_filename),
        os.path.join(TEMPLATE_DIR, template_filename)
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            return path
    
    # If no template found, use the fallback
    fallback_path = os.path.join(TEMPLATE_DIR, "fallback_template.docx")
    
    # Create a simple fallback template if it doesn't exist
    if not os.path.exists(fallback_path):
        create_fallback_template(fallback_path)
    
    return fallback_path

def create_fallback_template(path):
    """Create a simple fallback template if no specific template exists"""
    doc = Document()
    doc.add_heading('Legal Dispute Notification', 0)
    doc.add_paragraph('Date: {{date}}')
    doc.add_paragraph('To: {{recipient_name}}')
    doc.add_paragraph('From: {{full_name}}')
    doc.add_paragraph('Subject: OFFICIAL LEGAL DISPUTE NOTIFICATION')
    doc.add_paragraph()
    doc.add_paragraph('Dear Sir/Madam,')
    doc.add_paragraph()
    p = doc.add_paragraph('Under the authority of the ')
    p.add_run('{{legislation}}').bold = True
    p.add_run(', I am writing to formally dispute the following matter:')
    doc.add_paragraph()
    doc.add_paragraph('{{issue_description}}')
    doc.add_paragraph()
    doc.add_paragraph('{{additional_notes}}')
    doc.add_paragraph()
    doc.add_paragraph('I request that this matter be resolved in accordance with the applicable legislation. Please respond to this notice within 14 days of receipt.')
    doc.add_paragraph()
    doc.add_paragraph(f'If this matter cannot be resolved directly, I reserve the right to escalate this dispute to the {{authority}} or other appropriate legal channels.')
    doc.add_paragraph()
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph()
    doc.add_paragraph('____________________')
    doc.add_paragraph('{{full_name}}')
    doc.add_paragraph()
    doc.add_paragraph('This document was prepared with assistance from SmartDispute.ai, an AI-powered legal document service. The content is provided for informational purposes only and does not constitute legal advice.')
    
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(path), exist_ok=True)
    
    # Save the document
    doc.save(path)
    
    return path

def generate_document(user_data):
    """Generate a document using the appropriate template and user data"""
    province = user_data.get('province')
    dispute_type = user_data.get('dispute_type')
    
    # Get the template
    template_path = get_template_path(province, dispute_type)
    
    # Get legislation and authority information
    legislation = LEGISLATION_MAP.get(province, {}).get(dispute_type, "applicable provincial legislation")
    authority = AUTHORITY_MAP.get(province, {}).get(dispute_type, "relevant provincial authority")
    
    # Create a unique filename
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = str(uuid.uuid4())[:8]
    docx_filename = f"dispute_{province}_{dispute_type}_{timestamp}_{unique_id}.docx"
    pdf_filename = docx_filename.replace(".docx", ".pdf")
    
    docx_path = os.path.join(GENERATED_DIR, docx_filename)
    pdf_path = os.path.join(GENERATED_DIR, pdf_filename)
    
    # Prepare context for template
    context = {
        'date': datetime.datetime.now().strftime("%B %d, %Y"),
        'legislation': legislation,
        'authority': authority,
        **user_data
    }
    
    # Render the template
    try:
        # Try to use DocxTemplate first
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(docx_path)
    except Exception as e:
        # Fallback to creating a new document
        print(f"Error using template: {e}")
        create_doc_manually(docx_path, context)
    
    # Convert to PDF
    try:
        pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        # If conversion fails, just return the DOCX file
        return docx_filename
    
    return pdf_filename

def create_doc_manually(output_path, context):
    """Create a document manually if template rendering fails"""
    doc = Document()
    
    doc.add_heading('Legal Dispute Notification', 0)
    doc.add_paragraph(f"Date: {context.get('date')}")
    
    if context.get('recipient_name'):
        doc.add_paragraph(f"To: {context.get('recipient_name')}")
    
    doc.add_paragraph(f"From: {context.get('full_name')}")
    if context.get('email'):
        doc.add_paragraph(f"Email: {context.get('email')}")
    if context.get('phone'):
        doc.add_paragraph(f"Phone: {context.get('phone')}")
    
    doc.add_paragraph('Subject: OFFICIAL LEGAL DISPUTE NOTIFICATION')
    doc.add_paragraph('')
    doc.add_paragraph('Dear Sir/Madam,')
    doc.add_paragraph('')
    
    p = doc.add_paragraph('Under the authority of the ')
    p.add_run(context.get('legislation', 'applicable provincial legislation')).bold = True
    p.add_run(', I am writing to formally dispute the following matter:')
    
    doc.add_paragraph('')
    doc.add_paragraph(context.get('issue_description', 'No description provided.'))
    doc.add_paragraph('')
    
    if context.get('additional_notes'):
        doc.add_paragraph('Additional Information:')
        doc.add_paragraph(context.get('additional_notes'))
        doc.add_paragraph('')
    
    doc.add_paragraph('I request that this matter be resolved in accordance with the applicable legislation. Please respond to this notice within 14 days of receipt.')
    doc.add_paragraph('')
    
    authority = context.get('authority', 'relevant provincial authority')
    doc.add_paragraph(f'If this matter cannot be resolved directly, I reserve the right to escalate this dispute to the {authority} or other appropriate legal channels.')
    
    doc.add_paragraph('')
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph('')
    doc.add_paragraph('____________________')
    doc.add_paragraph(context.get('full_name', ''))
    doc.add_paragraph('')
    
    doc.add_paragraph('This document was prepared with assistance from SmartDispute.ai, an AI-powered legal document service. The content is provided for informational purposes only and does not constitute legal advice.')
    
    # Save the document
    doc.save(output_path)
    
    return output_path

# Routes
@app.route('/')
def index():
    """Display the province and dispute type selection page"""
    return render_template('index.html', provinces=LEGISLATION_MAP.keys())

@app.route('/api/dispute-types/<province>')
def get_dispute_types(province):
    """Get dispute types for a specific province"""
    if province in LEGISLATION_MAP:
        return jsonify(list(LEGISLATION_MAP[province].keys()))
    return jsonify([])

@app.route('/form')
def form():
    """Display the form for a specific province and dispute type"""
    province = request.args.get('province')
    dispute_type = request.args.get('dispute_type')
    
    if not province or not dispute_type:
        return redirect(url_for('index'))
    
    legislation = LEGISLATION_MAP.get(province, {}).get(dispute_type)
    authority = AUTHORITY_MAP.get(province, {}).get(dispute_type)
    
    return render_template('form.html', 
                           province=province,
                           dispute_type=dispute_type,
                           legislation=legislation,
                           authority=authority)

@app.route('/generate', methods=['POST'])
def generate():
    """Process form submission and generate document"""
    user_data = {
        'province': request.form.get('province'),
        'dispute_type': request.form.get('dispute_type'),
        'full_name': request.form.get('full_name'),
        'email': request.form.get('email'),
        'phone': request.form.get('phone'),
        'recipient_name': request.form.get('recipient_name'),
        'recipient_address': request.form.get('recipient_address'),
        'issue_description': request.form.get('issue_description'),
        'additional_notes': request.form.get('additional_notes')
    }
    
    # Generate document
    filename = generate_document(user_data)
    
    # Send confirmation email if email is provided
    if user_data.get('email') and os.environ.get('EMAIL_ENABLED') == 'true':
        try:
            send_confirmation_email(user_data, filename)
        except Exception as e:
            print(f"Error sending email: {e}")
    
    return redirect(url_for('success', file=filename))

@app.route('/success')
def success():
    """Display success page with download link"""
    filename = request.args.get('file')
    if not filename:
        return redirect(url_for('index'))
    
    return render_template('success.html', filename=filename)

@app.route('/download/<filename>')
def download(filename):
    """Download the generated document"""
    return send_from_directory(GENERATED_DIR, filename)

@app.route('/api/legislation-map')
def get_legislation_map():
    """API endpoint to get legislation map"""
    return jsonify(LEGISLATION_MAP)

def send_confirmation_email(user_data, filename):
    """Send confirmation email with document attachment"""
    try:
        # Use subprocess to call the Node.js email integration
        import subprocess
        
        # Prepare data to pass to the Node.js script
        import json
        user_data_json = json.dumps(user_data)
        
        # Run the Node.js script with the user data and filename
        cmd = [
            'node', 
            '-e', 
            f"""
            const integration = require('./flask_email_integration.cjs');
            const userData = JSON.parse('{user_data_json}');
            const filename = '{filename}';
            
            (async () => {{
                try {{
                    const result = await integration.sendGeneratedDocument(userData, filename);
                    console.log(JSON.stringify(result));
                    process.exit(result.success ? 0 : 1);
                }} catch (error) {{
                    console.error('Error:', error.message);
                    process.exit(1);
                }}
            }})();
            """
        ]
        
        # Run the command and capture output
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode != 0:
            print(f"Error sending email: {result.stderr}")
            return False
            
        print(f"Email sent successfully: {result.stdout}")
        return True
        
    except Exception as e:
        print(f"Exception sending email: {str(e)}")
        return False

if __name__ == '__main__':
    # Create a simple template if not exists
    fallback_path = os.path.join(TEMPLATE_DIR, "fallback_template.docx")
    if not os.path.exists(fallback_path):
        create_fallback_template(fallback_path)
    
    # Use a different port to avoid conflicts    
    port = int(os.environ.get('FLASK_PORT', 5050))
    app.run(host='0.0.0.0', port=port, debug=True)