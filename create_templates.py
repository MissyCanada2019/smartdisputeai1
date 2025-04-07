from docx import Document
import os

# Define the provinces and dispute types
provinces = ['ON', 'BC', 'AB', 'QC']
dispute_types = ['landlord_tenant', 'credit', 'cas']

TEMPLATE_FOLDER = 'templates/docs'

# Create the template folder if it doesn't exist
if not os.path.exists(TEMPLATE_FOLDER):
    os.makedirs(TEMPLATE_FOLDER)

# Create a template for each combination
for province in provinces:
    for dispute_type in dispute_types:
        # Template file path
        template_path = os.path.join(TEMPLATE_FOLDER, f"{dispute_type}_{province}.docx")
        
        # Skip if the template already exists
        if os.path.exists(template_path):
            print(f"Template already exists: {template_path}")
            continue
        
        # Create a new document
        doc = Document()
        
        # Add content to the document
        doc.add_heading(f'SmartDispute.ai - {dispute_type.replace("_", " ").title()} Dispute for {province}', 0)
        
        doc.add_paragraph('Recipient:')
        doc.add_paragraph('{{ recipient_name }}')
        doc.add_paragraph('{{ recipient_address }}')
        
        doc.add_paragraph('')
        doc.add_paragraph('Sender:')
        doc.add_paragraph('{{ name }}')
        doc.add_paragraph('{{ email }}')
        
        doc.add_paragraph('')
        doc.add_paragraph('Date: {{ now() }}')
        
        doc.add_paragraph('')
        doc.add_heading('RE: FORMAL DISPUTE NOTICE', 1)
        
        doc.add_paragraph('')
        doc.add_paragraph('Dear Sir/Madam,')
        
        doc.add_paragraph('')
        if dispute_type == 'landlord_tenant':
            if province == 'ON':
                doc.add_paragraph('I am writing in accordance with the Residential Tenancies Act, 2006 to formally dispute the following issue:')
            elif province == 'BC':
                doc.add_paragraph('I am writing in accordance with the Residential Tenancy Act to formally dispute the following issue:')
            elif province == 'AB':
                doc.add_paragraph('I am writing in accordance with the Residential Tenancies Act to formally dispute the following issue:')
            elif province == 'QC':
                doc.add_paragraph('I am writing in accordance with the Civil Code of Québec (articles 1851-1978) to formally dispute the following issue:')
        elif dispute_type == 'credit':
            if province == 'ON':
                doc.add_paragraph('I am writing in accordance with the Consumer Reporting Act to formally dispute the following issue:')
            elif province == 'BC':
                doc.add_paragraph('I am writing in accordance with the Business Practices and Consumer Protection Act to formally dispute the following issue:')
            elif province == 'AB':
                doc.add_paragraph('I am writing in accordance with the Consumer Protection Act to formally dispute the following issue:')
            elif province == 'QC':
                doc.add_paragraph('I am writing in accordance with the Consumer Protection Act to formally dispute the following issue:')
        elif dispute_type == 'cas':
            if province == 'ON':
                doc.add_paragraph('I am writing in accordance with the Child, Youth and Family Services Act, 2017 to formally address the following issue:')
            elif province == 'BC':
                doc.add_paragraph('I am writing in accordance with the Child, Family and Community Service Act to formally address the following issue:')
            elif province == 'AB':
                doc.add_paragraph('I am writing in accordance with the Child, Youth and Family Enhancement Act to formally address the following issue:')
            elif province == 'QC':
                doc.add_paragraph('I am writing in accordance with the Youth Protection Act to formally address the following issue:')
        
        doc.add_paragraph('')
        doc.add_paragraph('{{ description }}')
        
        doc.add_paragraph('')
        doc.add_paragraph('I request that this matter be addressed and resolved promptly according to applicable laws. Please respond within 14 days of receipt of this letter.')
        
        doc.add_paragraph('')
        if dispute_type == 'landlord_tenant':
            if province == 'ON':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Landlord and Tenant Board.')
            elif province == 'BC':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Residential Tenancy Branch.')
            elif province == 'AB':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Residential Tenancy Dispute Resolution Service.')
            elif province == 'QC':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Tribunal administratif du logement.')
        elif dispute_type == 'credit':
            if province == 'ON':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Financial Services Regulatory Authority of Ontario.')
            elif province == 'BC':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with Consumer Protection BC.')
            elif province == 'AB':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with Service Alberta, Consumer Investigations Unit.')
            elif province == 'QC':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Office de la protection du consommateur.')
        elif dispute_type == 'cas':
            if province == 'ON':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Ontario Association of Children\'s Aid Societies or seek appropriate legal counsel.')
            elif province == 'BC':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Ministry of Children and Family Development or seek appropriate legal counsel.')
            elif province == 'AB':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with Alberta Children\'s Services or seek appropriate legal counsel.')
            elif province == 'QC':
                doc.add_paragraph('If this matter cannot be resolved directly, I reserve the right to file a complaint with the Direction de la protection de la jeunesse or seek appropriate legal counsel.')
        
        doc.add_paragraph('')
        doc.add_paragraph('Sincerely,')
        doc.add_paragraph('')
        doc.add_paragraph('{{ name }}')
        
        doc.add_paragraph('')
        doc.add_paragraph('This document was prepared with assistance from SmartDispute.ai, an AI-powered legal document service. The content is provided for informational purposes only and does not constitute legal advice.')
        
        # Save the document
        doc.save(template_path)
        print(f"Created template: {template_path}")

print("Template creation completed.")