from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_credit_bc_template():
    """Create a specialized template for BC credit disputes"""
    document = Document()
    
    # Set up document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add heading
    heading = document.add_heading('CREDIT DISPUTE NOTIFICATION', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add recipient line
    document.add_paragraph("To Whom It May Concern,")
    document.add_paragraph()  # Blank line
    
    # Add introduction
    intro = document.add_paragraph()
    intro.add_run("My name is {{ name }}, and I am a resident of {{ province }}. I am contacting you to formally dispute information currently listed on my credit report. The disputed issue is described below:")
    
    # Add description placeholder
    document.add_paragraph()  # Blank line
    description_para = document.add_paragraph()
    description_para.add_run("{{ description }}")
    document.add_paragraph()  # Blank line
    
    # Add legal reference
    legal_para = document.add_paragraph()
    legal_para.add_run("Under the Business Practices and Consumer Protection Act (British Columbia), I am entitled to fair and accurate reporting. I request that this matter be reviewed immediately and the necessary corrections be made.")
    
    # Add documentation note
    document.add_paragraph()  # Blank line
    document.add_paragraph("Please find supporting documentation attached for your review.")
    
    # Add signature block
    document.add_paragraph()  # Blank line
    document.add_paragraph("Sincerely,")
    document.add_paragraph("{{ name }}")
    document.add_paragraph("{{ email }}")
    if "{{ address }}" != "":
        document.add_paragraph("{{ address }}")
    
    # Add date
    document.add_paragraph()  # Blank line
    date_para = document.add_paragraph()
    date_para.add_run("Date: {{ now() }}")
    
    # Ensure templates directory exists
    os.makedirs("templates/docs", exist_ok=True)
    
    # Save the document
    filename = "templates/docs/credit_BC.docx"
    document.save(filename)
    print(f"Created template: {filename}")
    
    return filename

def create_cas_ab_template():
    """Create a specialized template for Alberta CAS complaints"""
    document = Document()
    
    # Set up document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add heading
    heading = document.add_heading('FORMAL COMPLAINT – CHILD & FAMILY SERVICES', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add recipient line
    document.add_paragraph("To Whom It May Concern,")
    document.add_paragraph()  # Blank line
    
    # Add introduction
    intro = document.add_paragraph()
    intro.add_run("My name is {{ name }}, and I reside in {{ province }}. I am submitting a formal complaint regarding actions or inactions by Child and Family Services that I believe are inconsistent with the Child, Youth and Family Enhancement Act.")
    
    # Add description placeholder
    document.add_paragraph()  # Blank line
    document.add_paragraph("The concern I wish to raise is outlined below:")
    document.add_paragraph()  # Blank line
    description_para = document.add_paragraph()
    description_para.add_run("{{ description }}")
    document.add_paragraph()  # Blank line
    
    # Add urgency note
    legal_para = document.add_paragraph()
    legal_para.add_run("This matter requires urgent review to ensure that my rights and the well-being of the child(ren) involved are fully protected under Alberta legislation.")
    
    # Add evidence note
    document.add_paragraph()  # Blank line
    document.add_paragraph("I am prepared to provide further evidence or clarification as needed.")
    
    # Add signature block
    document.add_paragraph()  # Blank line
    document.add_paragraph("Sincerely,")
    document.add_paragraph("{{ name }}")
    document.add_paragraph("{{ email }}")
    if "{{ address }}" != "":
        document.add_paragraph("{{ address }}")
    
    # Add date
    document.add_paragraph()  # Blank line
    date_para = document.add_paragraph()
    date_para.add_run("Date: {{ now() }}")
    
    # Ensure templates directory exists
    os.makedirs("templates/docs", exist_ok=True)
    
    # Save the document
    filename = "templates/docs/cas_AB.docx"
    document.save(filename)
    print(f"Created template: {filename}")
    
    return filename

if __name__ == "__main__":
    bc_template = create_credit_bc_template()
    ab_template = create_cas_ab_template()
    print(f"Successfully created templates at {bc_template} and {ab_template}")