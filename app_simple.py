from flask import Flask, request, render_template, send_from_directory, redirect, url_for
from docxtpl import DocxTemplate
from docx import Document
import os
import uuid
import json
import datetime
import subprocess
from flask import jsonify, session
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'smartdispute_ai_development_key')

def get_legal_act(dispute_type, province):
    """Get the appropriate legal act based on dispute type and province"""
    legal_acts = {
        'landlord_tenant': {
            'ON': 'Residential Tenancies Act',
            'BC': 'Residential Tenancy Act',
            'AB': 'Residential Tenancies Act',
            'QC': 'Civil Code of Québec (sections governing residential leases)',
            'MB': 'Residential Tenancies Act',
            'SK': 'Residential Tenancies Act',
            'NS': 'Residential Tenancies Act',
            'NB': 'Residential Tenancies Act',
            'NL': 'Residential Tenancies Act',
            'PE': 'Rental of Residential Property Act',
            'NT': 'Residential Tenancies Act',
            'NU': 'Residential Tenancies Act',
            'YT': 'Residential Landlord and Tenant Act'
        },
        'landlord_cease_desist': {
            'ON': 'Residential Tenancies Act',
            'BC': 'Residential Tenancy Act',
            'AB': 'Residential Tenancies Act',
            'QC': 'Civil Code of Québec (sections governing residential leases)',
            'MB': 'Residential Tenancies Act',
            'SK': 'Residential Tenancies Act',
            'NS': 'Residential Tenancies Act',
            'NB': 'Residential Tenancies Act',
            'NL': 'Residential Tenancies Act',
            'PE': 'Rental of Residential Property Act',
            'NT': 'Residential Tenancies Act',
            'NU': 'Residential Tenancies Act',
            'YT': 'Residential Landlord and Tenant Act'
        },
        'repair_notice': {
            'ON': 'Residential Tenancies Act',
            'BC': 'Residential Tenancy Act',
            'AB': 'Residential Tenancies Act',
            'QC': 'Civil Code of Québec (sections governing residential leases)',
            'MB': 'Residential Tenancies Act',
            'SK': 'Residential Tenancies Act',
            'NS': 'Residential Tenancies Act',
            'NB': 'Residential Tenancies Act',
            'NL': 'Residential Tenancies Act',
            'PE': 'Rental of Residential Property Act',
            'NT': 'Residential Tenancies Act',
            'NU': 'Residential Tenancies Act',
            'YT': 'Residential Landlord and Tenant Act'
        },
        'intent_to_vacate': {
            'ON': 'Residential Tenancies Act',
            'BC': 'Residential Tenancy Act',
            'AB': 'Residential Tenancies Act',
            'QC': 'Civil Code of Québec (sections governing residential leases)',
            'MB': 'Residential Tenancies Act',
            'SK': 'Residential Tenancies Act',
            'NS': 'Residential Tenancies Act',
            'NB': 'Residential Tenancies Act',
            'NL': 'Residential Tenancies Act',
            'PE': 'Rental of Residential Property Act',
            'NT': 'Residential Tenancies Act',
            'NU': 'Residential Tenancies Act',
            'YT': 'Residential Landlord and Tenant Act'
        },
        'termination_notice': {
            'ON': 'Residential Tenancies Act',
            'BC': 'Residential Tenancy Act',
            'AB': 'Residential Tenancies Act',
            'QC': 'Civil Code of Québec (sections governing residential leases)',
            'MB': 'Residential Tenancies Act',
            'SK': 'Residential Tenancies Act',
            'NS': 'Residential Tenancies Act',
            'NB': 'Residential Tenancies Act',
            'NL': 'Residential Tenancies Act',
            'PE': 'Rental of Residential Property Act',
            'NT': 'Residential Tenancies Act',
            'NU': 'Residential Tenancies Act',
            'YT': 'Residential Landlord and Tenant Act'
        },
        'sublease_agreement': {
            'ON': 'Residential Tenancies Act',
            'BC': 'Residential Tenancy Act',
            'AB': 'Residential Tenancies Act',
            'QC': 'Civil Code of Québec (sections governing residential leases)',
            'MB': 'Residential Tenancies Act',
            'SK': 'Residential Tenancies Act',
            'NS': 'Residential Tenancies Act',
            'NB': 'Residential Tenancies Act',
            'NL': 'Residential Tenancies Act',
            'PE': 'Rental of Residential Property Act',
            'NT': 'Residential Tenancies Act',
            'NU': 'Residential Tenancies Act',
            'YT': 'Residential Landlord and Tenant Act'
        },
        'credit': {
            'ON': 'Consumer Reporting Act',
            'BC': 'Business Practices and Consumer Protection Act',
            'AB': 'Consumer Protection Act',
            'QC': 'Act Respecting the Protection of Personal Information in the Private Sector',
            'MB': 'Consumer Protection Act',
            'SK': 'Consumer Protection and Business Practices Act',
            'NS': 'Consumer Reporting Act',
            'NB': 'Consumer Reporting Act',
            'NL': 'Consumer Protection and Business Practices Act',
            'PE': 'Consumer Reporting Act',
            'NT': 'Consumer Protection Act',
            'NU': 'Consumer Protection Act',
            'YT': 'Consumer Protection Act'
        },
        'cas': {
            'ON': 'Child, Youth and Family Services Act',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act (Loi sur la protection de la jeunesse)',
            'MB': 'Child and Family Services Act',
            'SK': 'Child and Family Services Act',
            'NS': 'Children and Family Services Act',
            'NB': 'Family Services Act',
            'NL': 'Children, Youth and Families Act',
            'PE': 'Child Protection Act',
            'NT': 'Child and Family Services Act',
            'NU': 'Child and Family Services Act',
            'YT': 'Child and Family Services Act'
        },
        'cas_cease_desist': {
            'ON': 'Child, Youth and Family Services Act',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act (Loi sur la protection de la jeunesse)',
            'MB': 'Child and Family Services Act',
            'SK': 'Child and Family Services Act',
            'NS': 'Children and Family Services Act',
            'NB': 'Family Services Act',
            'NL': 'Children, Youth and Families Act',
            'PE': 'Child Protection Act',
            'NT': 'Child and Family Services Act',
            'NU': 'Child and Family Services Act',
            'YT': 'Child and Family Services Act'
        },
        'cas_worker_reassignment': {
            'ON': 'Child, Youth and Family Services Act, 2017',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act',
            'MB': 'Child and Family Services Act (Manitoba)',
            'SK': 'The Child and Family Services Act',
            'NS': 'Children and Family Services Act (Nova Scotia)',
            'NB': 'Family Services Act (New Brunswick)',
            'NL': 'Children, Youth and Families Act (Newfoundland and Labrador)',
            'PE': 'Child Protection Act (Prince Edward Island)',
            'NT': 'Child and Family Services Act (Northwest Territories)',
            'YT': 'Child and Family Services Act (Yukon)',
            'NU': 'Child and Family Services Act (Nunavut)'
        },
        'cas_answer_plan': {
            'ON': 'Child, Youth and Family Services Act, 2017',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act',
            'MB': 'Child and Family Services Act (Manitoba)',
            'SK': 'The Child and Family Services Act',
            'NS': 'Children and Family Services Act (Nova Scotia)',
            'NB': 'Family Services Act (New Brunswick)',
            'NL': 'Children, Youth and Families Act (Newfoundland and Labrador)',
            'PE': 'Child Protection Act (Prince Edward Island)',
            'NT': 'Child and Family Services Act (Northwest Territories)',
            'YT': 'Child and Family Services Act (Yukon)',
            'NU': 'Child and Family Services Act (Nunavut)'
        },
        'cas_records_request': {
            'ON': 'Child, Youth and Family Services Act, 2017',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act',
            'MB': 'Child and Family Services Act (Manitoba)',
            'SK': 'The Child and Family Services Act',
            'NS': 'Children and Family Services Act (Nova Scotia)',
            'NB': 'Family Services Act (New Brunswick)',
            'NL': 'Children, Youth and Families Act (Newfoundland and Labrador)',
            'PE': 'Child Protection Act (Prince Edward Island)',
            'NT': 'Child and Family Services Act (Northwest Territories)',
            'YT': 'Child and Family Services Act (Yukon)',
            'NU': 'Child and Family Services Act (Nunavut)'
        },
        'cas_appeal': {
            'ON': 'Child, Youth and Family Services Act, 2017',
            'BC': 'Child, Family and Community Service Act',
            'AB': 'Child, Youth and Family Enhancement Act',
            'QC': 'Youth Protection Act',
            'MB': 'Child and Family Services Act (Manitoba)',
            'SK': 'The Child and Family Services Act',
            'NS': 'Children and Family Services Act (Nova Scotia)',
            'NB': 'Family Services Act (New Brunswick)',
            'NL': 'Children, Youth and Families Act (Newfoundland and Labrador)',
            'PE': 'Child Protection Act (Prince Edward Island)',
            'NT': 'Child and Family Services Act (Northwest Territories)',
            'YT': 'Child and Family Services Act (Yukon)',
            'NU': 'Child and Family Services Act (Nunavut)'
        }
    }
    
    return legal_acts.get(dispute_type, {}).get(province, 'applicable provincial legislation')

def get_agency_name(province):
    """Get the appropriate child services agency name based on province code"""
    agency_map = {
        "ON": "Children's Aid Society",
        "BC": "Ministry of Children and Family Development",
        "AB": "Children's Services",
        "QC": "Direction de la protection de la jeunesse (DPJ)",
        "MB": "Child and Family Services",
        "SK": "Ministry of Social Services",
        "NS": "Department of Community Services",
        "NB": "Department of Social Development",
        "NL": "Department of Children, Seniors and Social Development",
        "PE": "Child Protection Services",
        "NT": "Child and Family Services",
        "NU": "Department of Family Services",
        "YT": "Family and Children's Services"
    }
    
    return agency_map.get(province, "Child Protection Agency")
UPLOAD_FOLDER = 'generated'
TEMPLATE_FOLDER = 'templates/docs'

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

if not os.path.exists(TEMPLATE_FOLDER):
    os.makedirs(TEMPLATE_FOLDER)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/form')
def dynamic_form():
    province = request.args.get('province')
    dispute_type = request.args.get('type')
    return render_template('form.html', province=province, dispute_type=dispute_type)

@app.route('/generate', methods=['POST'])
def generate():
    name = request.form['name']
    email = request.form['email']
    description = request.form['description']
    province = request.form['province']
    dispute_type = request.form['dispute_type']
    
    # Get optional form fields with defaults
    recipient_name = request.form.get('recipient_name', 'To Whom It May Concern')
    recipient_address = request.form.get('recipient_address', '')
    user_address = request.form.get('address', '')
    
    # Language preference for Quebec (default to English)
    language = request.form.get('language', 'EN')
    
    # Initialize all possible form fields to avoid unbound variable errors
    # Cease and desist fields
    prior_incidents = request.form.get('prior_incidents', '')
    requested_actions = request.form.get('requested_actions', '')
    consequences = request.form.get('consequences', '')
    
    # Worker reassignment fields
    current_worker_name = request.form.get('current_worker_name', '')
    description_of_concerns = request.form.get('description_of_concerns', '')
    previous_attempts = request.form.get('previous_attempts', '')
    
    # Repair notice fields
    repair_issues = request.form.get('repair_issues', '')
    prior_notifications = request.form.get('prior_notifications', '')
    
    # Intent to vacate fields
    move_out_date = request.form.get('move_out_date', '')
    tenancy_details = request.form.get('tenancy_details', '')
    
    # Termination notice fields
    termination_date = request.form.get('termination_date', '')
    termination_reason = request.form.get('termination_reason', '')
    
    # Sublease agreement fields
    subtenant_name = request.form.get('subtenant_name', '')
    property_address = request.form.get('property_address', '')
    sublease_start = request.form.get('sublease_start', '')
    sublease_end = request.form.get('sublease_end', '')
    monthly_rent = request.form.get('monthly_rent', '')
    security_deposit = request.form.get('security_deposit', '')
    utilities_included = request.form.get('utilities_included', '')
    
    # Get template filename based on dispute type and province
    filename = f"{dispute_type}_{province}.docx"
    
    # Special handling for Quebec French templates
    if province == 'QC' and language == 'FR' and dispute_type == 'cas_cease_desist':
        filename = f"{dispute_type}_{province}_FR.docx"
    
    template_path = os.path.join(TEMPLATE_FOLDER, filename)

    # Select appropriate template creation function based on dispute type
    if not os.path.exists(template_path):
        # Landlord tenant templates
        if dispute_type == 'landlord_cease_desist':
            try:
                from create_landlord_cease_desist_template import create_landlord_cease_desist_template
                create_landlord_cease_desist_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating landlord cease & desist template: {e}")
                create_basic_template(template_path)
        
        # Tenant notice templates (repair, vacate, termination)
        elif dispute_type in ['repair_notice', 'intent_to_vacate', 'termination_notice']:
            try:
                from create_tenant_notice_templates import (
                    create_repair_notice_template, 
                    create_intent_to_vacate_template, 
                    create_termination_notice_template,
                )
                if dispute_type == 'repair_notice':
                    create_repair_notice_template(province)
                elif dispute_type == 'intent_to_vacate':
                    create_intent_to_vacate_template(province)
                elif dispute_type == 'termination_notice':
                    create_termination_notice_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating {dispute_type} template: {e}")
                create_basic_template(template_path)
                
        # Sublease agreement template
        elif dispute_type == 'sublease_agreement':
            try:
                from create_sublease_agreement_template import create_sublease_agreement_template, create_french_quebec_template
                if province == 'QC' and language == 'FR':
                    create_french_quebec_template()
                else:
                    create_sublease_agreement_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating sublease agreement template: {e}")
                create_basic_template(template_path)
        
        # CAS templates
        elif dispute_type == 'cas_cease_desist':
            try:
                from create_cas_cease_desist_template import create_cas_cease_desist_template, create_french_quebec_template
                if province == 'QC' and language == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_cease_desist_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS cease & desist template: {e}")
                create_basic_template(template_path)
                
        elif dispute_type == 'cas_worker_reassignment':
            try:
                from create_cas_worker_reassignment_template import create_cas_worker_reassignment_template, create_french_quebec_template
                if province == 'QC' and language == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_worker_reassignment_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS worker reassignment template: {e}")
                create_basic_template(template_path)
                
        elif dispute_type == 'cas_answer_plan':
            try:
                from create_cas_answer_plan_template import create_cas_answer_plan_template, create_french_quebec_template
                if province == 'QC' and language == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_answer_plan_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS answer and plan template: {e}")
                create_basic_template(template_path)
                
        elif dispute_type == 'cas_records_request':
            try:
                from create_cas_records_request_template import create_cas_records_request_template, create_french_quebec_template
                if province == 'QC' and language == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_records_request_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS records request template: {e}")
                create_basic_template(template_path)
                
        elif dispute_type == 'cas_appeal':
            try:
                from create_cas_appeal_template import create_cas_appeal_template, create_french_quebec_template
                if province == 'QC' and language == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_appeal_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS appeal template: {e}")
                create_basic_template(template_path)
                
        else:
            # Fall back to basic template
            create_basic_template(template_path)

    # Get legal act based on dispute type and province
    legal_act = get_legal_act(dispute_type, province)
    
    # Get agency name for CAS templates
    agency_name = get_agency_name(province) if dispute_type.startswith('cas') else ''

    # Get payment method
    payment_method = request.form.get('payment_method', '')
    
    # Check if payment is required
    premium_templates = [
        'cas_worker_reassignment', 
        'cas_answer_plan', 
        'cas_records_request', 
        'cas_appeal',
        'sublease_agreement'
    ]
    requires_payment = dispute_type in premium_templates
    
    # If payment is required but payment processing is not completed
    if requires_payment and not request.args.get('payment_completed'):
        # Store form data in session
        session_data = {
            # Common fields
            'name': name,
            'email': email,
            'description': description,
            'province': province,
            'dispute_type': dispute_type,
            'recipient_name': recipient_name,
            'recipient_address': recipient_address,
            'user_address': user_address,
            'payment_method': payment_method,
            'language': request.form.get('language', 'EN'),
            
            # Cease and desist fields
            'prior_incidents': prior_incidents,
            'requested_actions': requested_actions,
            'consequences': consequences,
            
            # Worker reassignment fields
            'current_worker_name': current_worker_name,
            'description_of_concerns': description_of_concerns,
            'previous_attempts': previous_attempts,
            
            # Repair notice fields
            'repair_issues': repair_issues,
            'prior_notifications': prior_notifications,
            
            # Intent to vacate fields
            'move_out_date': move_out_date,
            'tenancy_details': tenancy_details,
            
            # Termination notice fields
            'termination_date': termination_date,
            'termination_reason': termination_reason,
            
            # Sublease agreement fields
            'subtenant_name': subtenant_name,
            'property_address': property_address,
            'sublease_start': sublease_start,
            'sublease_end': sublease_end,
            'monthly_rent': monthly_rent,
            'security_deposit': security_deposit,
            'utilities_included': utilities_included,
            
            # CAS answer plan fields
            'case_number': request.form.get('case_number', ''),
            'plan_date': request.form.get('plan_date', ''),
            'alternative_plan': request.form.get('alternative_plan', ''),
            
            # CAS records request fields
            'specific_records': request.form.get('specific_records', ''),
            'purpose': request.form.get('purpose', ''),
            
            # CAS appeal fields
            'decision_date': request.form.get('decision_date', ''),
            'decision_details': request.form.get('decision_details', ''),
            'appeal_grounds': request.form.get('appeal_grounds', ''),
            'requested_remedy': request.form.get('requested_remedy', '')
        }
        
        # Save session data to a temporary file
        session_id = str(uuid.uuid4())
        session_path = os.path.join(UPLOAD_FOLDER, f"session_{session_id}.json")
        with open(session_path, 'w') as f:
            json.dump(session_data, f)
        
        # Redirect to payment page
        return redirect(url_for('payment', 
                               session_id=session_id, 
                               amount='9.99', 
                               payment_method=payment_method))
    
    # Form data for all templates was already collected above
    
    # Get CAS answer plan fields if not already collected
    case_number = request.form.get('case_number', '')
    plan_date = request.form.get('plan_date', '')
    alternative_plan = request.form.get('alternative_plan', '')
    
    # Get CAS records request fields if not already collected
    specific_records = request.form.get('specific_records', '')
    request_purpose = request.form.get('purpose', '')
    
    # Get CAS appeal fields if not already collected
    decision_date = request.form.get('decision_date', '')
    decision_details = request.form.get('decision_details', '')
    appeal_grounds = request.form.get('appeal_grounds', '')
    requested_remedy = request.form.get('requested_remedy', '')
    
    doc = DocxTemplate(template_path)
    context = {
        # Common fields for all templates
        'name': name,
        'email': email,
        'description': description,
        'province': province,
        'recipient_name': recipient_name,
        'recipient_address': recipient_address,
        'address': user_address,
        'legal_act': legal_act,
        'agency_name': agency_name,
        
        # Cease and desist fields
        'prior_incidents': prior_incidents,
        'requested_actions': requested_actions, 
        'consequences': consequences,
        
        # Worker reassignment fields
        'current_worker_name': current_worker_name,
        'description_of_concerns': description_of_concerns,
        'previous_attempts': previous_attempts,
        
        # Repair notice fields
        'repair_issues': repair_issues,
        'prior_notifications': prior_notifications,
        
        # Intent to vacate fields
        'move_out_date': move_out_date,
        'tenancy_details': tenancy_details,
        
        # Termination notice fields
        'termination_date': termination_date,
        'termination_reason': termination_reason,
        
        # Sublease agreement fields
        'subtenant_name': subtenant_name,
        'property_address': property_address,
        'sublease_start': sublease_start,
        'sublease_end': sublease_end,
        'monthly_rent': monthly_rent,
        'security_deposit': security_deposit,
        'utilities_included': utilities_included,
        
        # CAS answer plan fields
        'case_number': case_number,
        'plan_date': plan_date,
        'alternative_plan': alternative_plan,
        
        # CAS records request fields
        'specific_records': specific_records,
        'request_purpose': request_purpose,
        
        # CAS appeal fields
        'decision_date': decision_date,
        'decision_details': decision_details,
        'appeal_grounds': appeal_grounds,
        'requested_remedy': requested_remedy,
        
        # Date function
        'now': lambda: datetime.datetime.now().strftime("%B %d, %Y")
    }

    doc.render(context)
    output_filename = f"{uuid.uuid4()}.docx"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)
    doc.save(output_path)

    return redirect(url_for('success', filename=output_filename))

def create_basic_template(path):
    """Create a basic template if none exists"""
    doc = Document()
    doc.add_heading('SmartDispute.ai Generated Document', 0)
    
    # Sender information
    doc.add_heading('Sender Information', level=1)
    doc.add_paragraph('Name: {{ name }}')
    doc.add_paragraph('Email: {{ email }}')
    if '{{ address }}':
        doc.add_paragraph('Address: {{ address }}')
    doc.add_paragraph('Date: {{ now() }}')
    
    # Recipient information
    doc.add_heading('Recipient Information', level=1)
    doc.add_paragraph('To: {{ recipient_name }}')
    if '{{ recipient_address }}':
        doc.add_paragraph('{{ recipient_address }}')
    
    # Legal information
    doc.add_heading('Legal References', level=1)
    doc.add_paragraph('Province: {{ province }}')
    if '{{ legal_act }}':
        doc.add_paragraph('Applicable Legislation: {{ legal_act }}')
    if '{{ agency_name }}':
        doc.add_paragraph('Agency: {{ agency_name }}')
    
    # Issue description
    doc.add_heading('Issue Description', level=1)
    doc.add_paragraph('{{ description }}')
    
    # Template type specific sections
    
    # Tenant notice templates
    if path.find('repair_notice') != -1:
        doc.add_heading('Repair Details', level=1)
        doc.add_paragraph('Issue to be Repaired: {{ description }}')
        doc.add_paragraph('Requested Timeline: Within 14 days of receipt of this notice, as required by law')
        doc.add_paragraph('Consequences of Non-Action: If these repairs are not addressed within the specified timeframe, I reserve the right to pursue all available legal remedies.')
    
    elif path.find('intent_to_vacate') != -1:
        doc.add_heading('Notice to Vacate', level=1)
        doc.add_paragraph('This letter serves as my official notice of intent to vacate the rental property.')
        doc.add_paragraph('I will be vacating the property on or before: (date to be filled in)')
        doc.add_paragraph('I am providing this notice in accordance with the required notice period under provincial law.')
    
    elif path.find('termination_notice') != -1:
        doc.add_heading('Lease Termination', level=1)
        doc.add_paragraph('This letter serves as my official notice to terminate my tenancy agreement.')
        doc.add_paragraph('Reason for Termination: {{ description }}')
        doc.add_paragraph('I am providing this notice in accordance with the required notice period under provincial law.')
    
    elif path.find('sublease_agreement') != -1:
        doc.add_heading('Sublease Terms', level=1)
        doc.add_paragraph('This document serves as a formal sublease agreement.')
        doc.add_paragraph('Original Tenant: {{ name }}')
        doc.add_paragraph('Subtenant: [NAME OF SUBTENANT]')
        doc.add_paragraph('Sublease Period: From [START DATE] to [END DATE]')
        doc.add_paragraph('Monthly Rent: $[AMOUNT]')
        doc.add_paragraph('Please review the terms and conditions outlined in this agreement.')
    
    # Cease and desist specific sections
    elif path.find('cease_desist') != -1:
        doc.add_heading('Prior Incidents', level=1)
        doc.add_paragraph('{{ prior_incidents }}')
        
        doc.add_heading('Requested Actions', level=1)
        doc.add_paragraph('{{ requested_actions }}')
        
        doc.add_heading('Consequences', level=1)
        doc.add_paragraph('{{ consequences }}')
    
    # Worker reassignment specific sections
    elif path.find('worker_reassignment') != -1:
        doc.add_heading('Current CAS Worker', level=1)
        doc.add_paragraph('{{ current_worker_name }}')
        
        doc.add_heading('Concerns', level=1)
        doc.add_paragraph('{{ description_of_concerns }}')
        
        doc.add_heading('Previous Attempts to Resolve', level=1)
        doc.add_paragraph('{{ previous_attempts }}')
    
    # CAS Answer and Plan of Care sections
    elif path.find('cas_answer_plan') != -1:
        doc.add_heading('Response to Allegations', level=1)
        doc.add_paragraph('I am responding to the allegations outlined in the Plan of Care dated [DATE].')
        doc.add_paragraph('My position regarding these allegations: {{ description }}')
        
        doc.add_heading('Proposed Plan of Care', level=1)
        doc.add_paragraph('I propose the following alternative plan of care:')
        doc.add_paragraph('1. [DETAILS TO BE FILLED IN]')
        doc.add_paragraph('2. [DETAILS TO BE FILLED IN]')
        doc.add_paragraph('3. [DETAILS TO BE FILLED IN]')
    
    # CAS Records Request sections
    elif path.find('cas_records_request') != -1:
        doc.add_heading('Records Requested', level=1)
        doc.add_paragraph('Pursuant to applicable privacy legislation, I formally request copies of all records relating to:')
        doc.add_paragraph('1. [SPECIFIC RECORDS REQUESTED]')
        doc.add_paragraph('Please provide these records within the timeframe specified by law (typically 30 days).')
        
        doc.add_heading('Basis for Request', level=1)
        doc.add_paragraph('{{ description }}')
    
    # CAS Appeal sections
    elif path.find('cas_appeal') != -1:
        doc.add_heading('Decision Being Appealed', level=1)
        doc.add_paragraph('I am formally appealing the decision dated [DATE] regarding:')
        doc.add_paragraph('{{ description }}')
        
        doc.add_heading('Grounds for Appeal', level=1)
        doc.add_paragraph('1. [GROUNDS FOR APPEAL]')
        doc.add_paragraph('2. [GROUNDS FOR APPEAL]')
        
        doc.add_heading('Requested Remedy', level=1)
        doc.add_paragraph('I request that the original decision be [overturned/modified] as follows:')
        doc.add_paragraph('[REQUESTED REMEDY]')
    
    # Signature
    doc.add_paragraph('\n\nSincerely,\n\n\n\n{{ name }}')
    
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(path), exist_ok=True)
    
    doc.save(path)
    print(f"Created basic template at {path}")

@app.route('/payment')
def payment():
    """Handle payment page for paid documents"""
    session_id = request.args.get('session_id')
    amount = request.args.get('amount', '9.99')
    payment_method = request.args.get('payment_method', '')
    
    return render_template('payment.html', 
                           session_id=session_id, 
                           amount=amount, 
                           payment_method=payment_method)

@app.route('/paypal-payment')
def paypal_payment():
    """Handle PayPal payment page"""
    return render_template('paypal_payment_page.html', paypal_client_id=os.environ.get('PAYPAL_CLIENT_ID', ''))

@app.route('/payment-success')
def payment_success():
    """Handle payment success page"""
    # Get order details from query parameters or use defaults
    order_id = request.args.get('order_id', 'ORDER-' + str(uuid.uuid4())[:8])
    amount = request.args.get('amount', '49.99')
    currency = request.args.get('currency', 'CAD')
    date = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    return render_template('payment_success.html', 
                           order_id=order_id,
                           amount=amount,
                           currency=currency,
                           date=date)

@app.route('/process-payment', methods=['POST'])
def process_payment():
    """Process the payment and generate the document"""
    session_id = request.form.get('session_id')
    amount = request.form.get('amount')
    payment_method = request.form.get('payment_method')
    
    # In a real implementation, we would integrate with PayPal or a credit card processor here
    # For now, we'll simulate a successful payment
    
    # Load the stored session data
    session_path = os.path.join(UPLOAD_FOLDER, f"session_{session_id}.json")
    if not os.path.exists(session_path):
        return "Session expired or invalid. Please try again.", 400
    
    with open(session_path, 'r') as f:
        session_data = json.load(f)
    
    # Get the necessary data from the session
    name = session_data.get('name')
    email = session_data.get('email')
    description = session_data.get('description')
    province = session_data.get('province')
    dispute_type = session_data.get('dispute_type')
    recipient_name = session_data.get('recipient_name')
    recipient_address = session_data.get('recipient_address')
    user_address = session_data.get('user_address')
    current_worker_name = session_data.get('current_worker_name')
    description_of_concerns = session_data.get('description_of_concerns')
    previous_attempts = session_data.get('previous_attempts', '')
    
    # Get template filename and path
    filename = f"{dispute_type}_{province}.docx"
    template_path = os.path.join(TEMPLATE_FOLDER, filename)
    
    # Handle Quebec French template
    if province == 'QC' and session_data.get('language') == 'FR':
        filename = f"{dispute_type}_{province}_FR.docx"
        template_path = os.path.join(TEMPLATE_FOLDER, filename)
    
    # Make sure the template exists
    if not os.path.exists(template_path):
        # Select appropriate template creation function based on dispute type
        if dispute_type == 'cas_worker_reassignment':
            try:
                from create_cas_worker_reassignment_template import create_cas_worker_reassignment_template, create_french_quebec_template
                
                if province == 'QC' and session_data.get('language') == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_worker_reassignment_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS worker reassignment template: {e}")
                create_basic_template(template_path)
        
        elif dispute_type == 'cas_answer_plan':
            try:
                from create_cas_answer_plan_template import create_cas_answer_plan_template, create_french_quebec_template
                if province == 'QC' and session_data.get('language') == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_answer_plan_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS answer and plan template: {e}")
                create_basic_template(template_path)
                
        elif dispute_type == 'cas_records_request':
            try:
                from create_cas_records_request_template import create_cas_records_request_template, create_french_quebec_template
                if province == 'QC' and session_data.get('language') == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_records_request_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS records request template: {e}")
                create_basic_template(template_path)
                
        elif dispute_type == 'cas_appeal':
            try:
                from create_cas_appeal_template import create_cas_appeal_template, create_french_quebec_template
                if province == 'QC' and session_data.get('language') == 'FR':
                    create_french_quebec_template()
                else:
                    create_cas_appeal_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating CAS appeal template: {e}")
                create_basic_template(template_path)
                
        elif dispute_type == 'sublease_agreement':
            try:
                from create_sublease_agreement_template import create_sublease_agreement_template, create_french_quebec_template
                if province == 'QC' and session_data.get('language') == 'FR':
                    create_french_quebec_template()
                else:
                    create_sublease_agreement_template(province)
            except (ImportError, Exception) as e:
                print(f"Error creating sublease agreement template: {e}")
                create_basic_template(template_path)
        
        else:
            # Fall back to basic template
            create_basic_template(template_path)
    
    # Get legal act and agency name
    legal_act = get_legal_act(dispute_type, province)
    agency_name = get_agency_name(province)
    
    # Get additional form data from session
    # CAS answer plan fields
    case_number = session_data.get('case_number', '')
    plan_date = session_data.get('plan_date', '')
    alternative_plan = session_data.get('alternative_plan', '')
    
    # CAS records request fields
    specific_records = session_data.get('specific_records', '')
    request_purpose = session_data.get('purpose', '')
    
    # CAS appeal fields
    decision_date = session_data.get('decision_date', '')
    decision_details = session_data.get('decision_details', '')
    appeal_grounds = session_data.get('appeal_grounds', '')
    requested_remedy = session_data.get('requested_remedy', '')
    
    # Sublease agreement fields
    subtenant_name = session_data.get('subtenant_name', '')
    property_address = session_data.get('property_address', '')
    sublease_start = session_data.get('sublease_start', '')
    sublease_end = session_data.get('sublease_end', '')
    monthly_rent = session_data.get('monthly_rent', '')
    security_deposit = session_data.get('security_deposit', '')
    utilities_included = session_data.get('utilities_included', '')
    
    # Create the document
    doc = DocxTemplate(template_path)
    context = {
        # Common fields for all templates
        'name': name,
        'email': email,
        'description': description,
        'province': province,
        'recipient_name': recipient_name,
        'recipient_address': recipient_address,
        'address': user_address,
        'legal_act': legal_act,
        'agency_name': agency_name,
        
        # Worker reassignment fields
        'current_worker_name': current_worker_name,
        'description_of_concerns': description_of_concerns,
        'previous_attempts': previous_attempts,
        
        # CAS answer plan fields
        'case_number': case_number,
        'plan_date': plan_date,
        'alternative_plan': alternative_plan,
        
        # CAS records request fields
        'specific_records': specific_records,
        'request_purpose': request_purpose,
        
        # CAS appeal fields
        'decision_date': decision_date,
        'decision_details': decision_details,
        'appeal_grounds': appeal_grounds,
        'requested_remedy': requested_remedy,
        
        # Sublease agreement fields
        'subtenant_name': subtenant_name,
        'property_address': property_address,
        'sublease_start': sublease_start,
        'sublease_end': sublease_end,
        'monthly_rent': monthly_rent,
        'security_deposit': security_deposit,
        'utilities_included': utilities_included,
        
        # Date function
        'now': lambda: datetime.datetime.now().strftime("%B %d, %Y")
    }
    
    # Render and save the document
    doc.render(context)
    output_filename = f"{uuid.uuid4()}.docx"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)
    doc.save(output_path)
    
    # Clean up the session file
    try:
        os.remove(session_path)
    except:
        pass
    
    # Redirect to success page
    return redirect(url_for('success', filename=output_filename, payment='completed'))

@app.route('/success')
def success():
    """Display success page with document download link and email option"""
    filename = request.args.get('filename')
    payment_completed = request.args.get('payment')
    
    # Get the path to the generated file
    full_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(full_path):
        return 'File not found', 404
    
    # Check if email functionality is available (environment variables set)
    has_email_config = bool(os.environ.get('GMAIL_USER') and os.environ.get('GMAIL_APP_PASS'))
    
    return render_template('success.html', 
                          filename=filename, 
                          payment_completed=payment_completed,
                          email_enabled=has_email_config)

@app.route('/generated/<filename>')
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)

@app.route('/send-email', methods=['POST'])
def send_email():
    """Send the generated document via email"""
    filename = request.form.get('filename')
    recipient_email = request.form.get('email')
    user_name = request.form.get('name', '')
    dispute_type = request.form.get('dispute_type', 'legal document')
    custom_message = request.form.get('message', '')
    
    # Check required fields
    if not filename or not recipient_email:
        return jsonify({'success': False, 'error': 'Missing required fields'}), 400
    
    # Check if email configuration is available
    if not os.environ.get('GMAIL_USER') or not os.environ.get('GMAIL_APP_PASS'):
        return jsonify({'success': False, 'error': 'Email service not configured'}), 500
    
    # Get the path to the generated file
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(file_path):
        return jsonify({'success': False, 'error': 'File not found'}), 404
    
    try:
        # For a real implementation, we would use node to call the EmailService
        # This is a placeholder implementation using Flask's built-in capabilities
        import subprocess
        
        # Construct command to call the Node.js email service
        cmd = [
            'node', '-e', 
            f'''
            require('./flask_email_integration').sendUserNotification({{
                userEmail: "{recipient_email}",
                userName: "{user_name}",
                documentPath: "{file_path}",
                disputeType: "{dispute_type}",
                customMessage: "{custom_message}"
            }}).then(result => console.log(JSON.stringify(result)));
            '''
        ]
        
        # Execute the Node.js command to send email
        process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate()
        
        if process.returncode != 0:
            print(f"Error calling email service: {stderr.decode('utf-8')}")
            return jsonify({'success': False, 'error': 'Failed to send email'}), 500
        
        # Parse output to get result
        result = stdout.decode('utf-8').strip()
        print(f"Email service result: {result}")
        
        return jsonify({'success': True, 'message': 'Email sent successfully'})
        
    except Exception as e:
        print(f"Exception in send_email: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Set up uploads folder for document analysis
ANALYSIS_UPLOAD_FOLDER = 'uploads/documents'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'rtf', 'jpg', 'jpeg', 'png'}

# Create uploads directory if it doesn't exist
if not os.path.exists(ANALYSIS_UPLOAD_FOLDER):
    os.makedirs(ANALYSIS_UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    """Check if the file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/document-analysis', methods=['GET'])
def document_analysis_form():
    """Display the document analysis form"""
    return render_template('document_analysis.html')

@app.route('/analyze-document', methods=['POST'])
def analyze_document():
    """Process document for AI analysis"""
    try:
        # Check if a file was uploaded
        if 'document' not in request.files:
            return jsonify({"success": False, "message": "No document file provided"})
            
        file = request.files['document']
        
        # Check if file was selected
        if file.filename == '':
            return jsonify({"success": False, "message": "No file selected"})
            
        # Check if the file extension is allowed
        if not allowed_file(file.filename):
            allowed_ext = ', '.join(ALLOWED_EXTENSIONS)
            return jsonify({"success": False, "message": f"File type not allowed. Supported formats: {allowed_ext}"})
        
        # Get form data
        analysis_type = request.form.get('analysis_type', 'basic')
        province = request.form.get('province', 'ON')
        dispute_type = request.form.get('dispute_type', 'landlord_tenant')
        document_type = request.form.get('document_type', 'notice')
        
        # Save the uploaded file
        filename = secure_filename(file.filename)
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        file_path = os.path.join(ANALYSIS_UPLOAD_FOLDER, unique_filename)
        file.save(file_path)
        
        # Store file path in session for later use
        session['analysis_document_path'] = file_path
        session['analysis_document_name'] = filename
        
        # Free tier users get basic analysis only
        is_premium = request.form.get('premium', 'false').lower() == 'true'
        basic_analysis_only = not is_premium
        
        # If free tier, show payment page before full analysis
        if not is_premium and analysis_type == 'comprehensive':
            return redirect(url_for('document_analysis_payment', 
                                   province=province, 
                                   dispute_type=dispute_type,
                                   document_type=document_type))
        
        # Run the analysis using Node.js script
        result = subprocess.run([
            'node', 'document-analysis-test.cjs',
            '--document', file_path,
            '--province', province,
            '--dispute_type', dispute_type,
            '--document_type', document_type,
            '--basic_only', 'true' if basic_analysis_only else 'false'
        ], capture_output=True, text=True)
        
        # If the analysis was successful
        if result.returncode == 0:
            try:
                # Try to parse the output as JSON
                analysis_result = json.loads(result.stdout)
                
                # Store analysis result in session
                session['analysis_result'] = analysis_result
                
                # Redirect to analysis results page
                return redirect(url_for('document_analysis_result'))
            except json.JSONDecodeError:
                print("Failed to parse analysis result as JSON")
                return jsonify({
                    "success": False, 
                    "message": "Failed to parse analysis result",
                    "output": result.stdout
                })
        else:
            print(f"Error analyzing document: {result.stderr}")
            # Check for API key errors
            if "API key" in result.stderr:
                return jsonify({
                    "success": False, 
                    "message": "AI analysis service unavailable. Please try again later."
                })
            return jsonify({"success": False, "message": "Failed to analyze document"})
            
    except Exception as e:
        print(f"Exception in analyze_document: {str(e)}")
        return jsonify({"success": False, "message": f"Error processing document: {str(e)}"})

@app.route('/document-analysis-payment')
def document_analysis_payment():
    """Display payment page for document analysis"""
    province = request.args.get('province', 'ON')
    dispute_type = request.args.get('dispute_type', 'landlord_tenant')
    document_type = request.args.get('document_type', 'notice')
    
    return render_template('document_analysis_payment.html', 
                          province=province,
                          dispute_type=dispute_type,
                          document_type=document_type)

@app.route('/document-analysis-result')
def document_analysis_result():
    """Display document analysis results"""
    # Get analysis result from session
    analysis_result = session.get('analysis_result', {})
    document_name = session.get('analysis_document_name', 'Uploaded document')
    
    if not analysis_result:
        return redirect(url_for('document_analysis_form'))
    
    return render_template('document_analysis_result.html', 
                          analysis=analysis_result,
                          document_name=document_name)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5050)