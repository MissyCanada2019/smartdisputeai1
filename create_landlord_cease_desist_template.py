from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def get_provincial_legislation(province_code):
    """Return the appropriate legislation name based on province code"""
    legislation_map = {
        "ON": "Residential Tenancies Act",
        "BC": "Residential Tenancy Act",
        "AB": "Residential Tenancies Act",
        "QC": "Civil Code of Québec (sections governing residential leases)",
        "MB": "Residential Tenancies Act",
        "SK": "Residential Tenancies Act",
        "NS": "Residential Tenancies Act",
        "NB": "Residential Tenancies Act",
        "NL": "Residential Tenancies Act",
        "PE": "Rental of Residential Property Act",
        "NT": "Residential Tenancies Act",
        "NU": "Residential Tenancies Act",
        "YT": "Residential Landlord and Tenant Act"
    }
    
    return legislation_map.get(province_code, "applicable residential tenancy legislation")

def create_landlord_cease_desist_template(province_code=None):
    """Create a specialized template for landlord cease and desist letters"""
    document = Document()
    
    # Set up document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Get province-specific legislation if provided
    legal_act = get_provincial_legislation(province_code) if province_code else "{{ legal_act }}"
    province = province_code if province_code else "{{ province }}"
    
    # Add sender info placeholders
    sender_info = document.add_paragraph()
    sender_info.add_run("{{ name }}\n")
    sender_info.add_run("{{ address }}\n")
    sender_info.add_run("{{ email }}")
    if province_code:
        sender_info.add_run(f"\n{province}")
    sender_info.add_run("\n{{ now() }}")
    
    # Add space after sender info
    document.add_paragraph()
    
    # Add recipient info placeholders
    recipient_info = document.add_paragraph()
    recipient_info.add_run("{{ recipient_name }}\n")
    recipient_info.add_run("{{ recipient_address }}")
    
    # Add space after recipient info
    document.add_paragraph()
    
    # Add subject line
    subject = document.add_paragraph()
    subject_run = subject.add_run("Re: Formal Cease and Desist Demand – Unlawful Interference and Harassment")
    subject_run.bold = True
    
    # Add greeting
    document.add_paragraph("To Whom It May Concern,")
    
    # Add main body
    main_body = document.add_paragraph()
    main_body.add_run("I am writing to formally demand that you immediately cease and desist all actions that constitute interference with my legal right to quiet enjoyment of the rental unit located at ")
    main_body.add_run("{{ address }}").italic = True
    main_body.add_run(".")
    
    # Add list of actions
    document.add_paragraph()
    document.add_paragraph("This includes, but is not limited to:")
    bullet_points = document.add_paragraph("- Unlawful entry or threats of entry without proper notice\n")
    bullet_points.add_run("- Intimidation, harassment, or verbal abuse\n")
    bullet_points.add_run("- Attempts to force eviction through coercion or retaliatory action")
    
    # Add legal reference
    document.add_paragraph()
    legal_para = document.add_paragraph()
    legal_para.add_run(f"Under the ")
    legal_act_run = legal_para.add_run(f"{legal_act}")
    legal_act_run.bold = True
    legal_para.add_run(f" applicable in ")
    province_run = legal_para.add_run(f"{province}")
    province_run.bold = True
    legal_para.add_run(", I am entitled to peaceful occupancy and protection from harassment or unreasonable interference by my landlord or their agents.")
    
    # Add warning
    document.add_paragraph()
    document.add_paragraph("Let this serve as your official notice: if these actions continue, I will have no choice but to pursue remedies through the appropriate housing tribunal or court in accordance with the laws of {{ province }}. This may include filing for compensation, injunctions, or further legal action.")
    
    # Add request for acknowledgment
    document.add_paragraph()
    document.add_paragraph("I request that you acknowledge receipt of this letter and confirm that you will comply with this cease and desist demand within 5 business days.")
    
    # Add signature
    document.add_paragraph()
    signature = document.add_paragraph("Sincerely,")
    signature.add_run("\n{{ name }}")
    signature.add_run("\n{{ email }}")
    
    # Ensure templates directory exists
    os.makedirs("templates/docs", exist_ok=True)
    
    # Save the document with province-specific filename if provided
    if province_code:
        filename = f"templates/docs/landlord_cease_desist_{province_code}.docx"
    else:
        filename = "templates/docs/landlord_cease_desist.docx"
    
    document.save(filename)
    print(f"Created template: {filename}")
    
    return filename

def create_all_provincial_templates():
    """Create landlord cease and desist templates for all provinces"""
    provinces = ["ON", "BC", "AB", "QC", "MB", "SK", "NS", "NB", "NL", "PE", "NT", "NU", "YT"]
    created_files = []
    
    for province in provinces:
        filename = create_landlord_cease_desist_template(province)
        created_files.append(filename)
    
    return created_files

if __name__ == "__main__":
    # Create generic template
    generic_template = create_landlord_cease_desist_template()
    print(f"Created generic template: {generic_template}")
    
    # Create province-specific templates
    provincial_templates = create_all_provincial_templates()
    print(f"Created {len(provincial_templates)} provincial templates.")