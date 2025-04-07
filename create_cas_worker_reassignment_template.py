from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def get_provincial_legislation(province_code):
    """Return the appropriate child protection legislation based on province code"""
    legislation_map = {
        "ON": "Child, Youth and Family Services Act, 2017",
        "BC": "Child, Family and Community Service Act",
        "AB": "Child, Youth and Family Enhancement Act",
        "QC": "Youth Protection Act",
        "MB": "Child and Family Services Act (Manitoba)",
        "SK": "The Child and Family Services Act",
        "NS": "Children and Family Services Act (Nova Scotia)",
        "NB": "Family Services Act (New Brunswick)",
        "NL": "Children, Youth and Families Act (Newfoundland and Labrador)",
        "PE": "Child Protection Act (Prince Edward Island)",
        "NT": "Child and Family Services Act (Northwest Territories)",
        "YT": "Child and Family Services Act (Yukon)",
        "NU": "Child and Family Services Act (Nunavut)"
    }
    
    return legislation_map.get(province_code, "applicable child protection legislation")

def get_agency_name(province_code):
    """Get the appropriate child services agency name based on province code"""
    agency_map = {
        "ON": "Children's Aid Society",
        "BC": "Ministry of Children and Family Development",
        "AB": "Children's Services",
        "QC": "Direction de la protection de la jeunesse (DPJ)",
        "MB": "Child and Family Services",
        "SK": "Ministry of Social Services",
        "NS": "Department of Community Services",
        "NB": "Department of Social Development",
        "NL": "Department of Children, Seniors and Social Development",
        "PE": "Child Protection Services",
        "NT": "Child and Family Services",
        "NU": "Department of Family Services",
        "YT": "Family and Children's Services"
    }
    
    return agency_map.get(province_code, "Child Protection Agency")

def create_french_quebec_template():
    """Create a specialized French template for Quebec DPJ"""
    template_path = "templates/docs/cas_worker_reassignment_QC_FR.docx"
    doc = Document()
    
    # Document properties
    doc.core_properties.author = "SmartDispute.ai"
    doc.core_properties.title = "Demande officielle de réaffectation du travailleur de la DPJ"
    
    # Header section
    paragraph = doc.add_paragraph("{{ name }}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("{{ address }}")
    doc.add_paragraph("{{ email }}")
    doc.add_paragraph("Date: {{ now() }}")
    doc.add_paragraph()
    
    # Recipient
    paragraph = doc.add_paragraph("{{ recipient_name }}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("{{ recipient_address }}")
    doc.add_paragraph()
    
    # Subject line
    paragraph = doc.add_paragraph("Objet: Demande officielle de réaffectation du travailleur de la DPJ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.runs[0].bold = True
    doc.add_paragraph()
    
    # Greeting
    doc.add_paragraph("À qui de droit,")
    doc.add_paragraph()
    
    # Body
    body = doc.add_paragraph("Je vous écris pour demander officiellement la réaffectation de notre intervenant actuel, ")
    body.add_run("{{ current_worker_name }}").italic = True
    body.add_run(", en vertu de la Loi sur la protection de la jeunesse au Québec.")
    
    doc.add_paragraph("Cette demande est faite de bonne foi en raison des préoccupations et des circonstances suivantes:")
    doc.add_paragraph()
    
    concerns = doc.add_paragraph("{{ description_of_concerns }}")
    concerns.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    doc.add_paragraph("Je crois que continuer à travailler avec l'intervenant actuellement assigné peut compromettre l'intégrité de notre dossier et n'est pas propice à une relation de travail respectueuse ou constructive. Il est dans l'intérêt de toutes les parties, en particulier de l'enfant (des enfants) concerné(s), qu'un nouvel intervenant soit assigné qui puisse aborder ce dossier avec impartialité et professionnalisme.")
    doc.add_paragraph()
    
    doc.add_paragraph("Veuillez confirmer la réception de cette demande et fournir une confirmation écrite des prochaines étapes dans les 5 jours ouvrables.")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Cordialement,")
    doc.add_paragraph()
    doc.add_paragraph("{{ name }}")
    doc.add_paragraph("{{ email }}")
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    
    # Save the document
    doc.save(template_path)
    print(f"Created template: {template_path}")
    
    return template_path

def create_cas_worker_reassignment_template(province_code=None):
    """Create a specialized template for CAS worker reassignment requests"""
    # Set appropriate template path based on province
    if province_code:
        template_path = f"templates/docs/cas_worker_reassignment_{province_code}.docx"
    else:
        template_path = "templates/docs/cas_worker_reassignment.docx"
        
    doc = Document()
    
    # Get province-specific legislation and agency name
    legislation = get_provincial_legislation(province_code) if province_code else "{{ legal_act }}"
    agency_name = get_agency_name(province_code) if province_code else "{{ agency_name }}"
    
    # Document properties
    doc.core_properties.author = "SmartDispute.ai"
    doc.core_properties.title = "Formal Request for CAS Worker Reassignment"
    
    # Header section
    paragraph = doc.add_paragraph("{{ name }}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("{{ address }}")
    doc.add_paragraph("{{ email }}")
    doc.add_paragraph("Date: {{ now() }}")
    doc.add_paragraph()
    
    # Recipient
    paragraph = doc.add_paragraph("{{ recipient_name }}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("{{ recipient_address }}")
    doc.add_paragraph()
    
    # Subject line
    paragraph = doc.add_paragraph(f"Re: Formal Request for Reassignment of CAS Worker")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.runs[0].bold = True
    doc.add_paragraph()
    
    # Greeting
    doc.add_paragraph("To Whom It May Concern,")
    doc.add_paragraph()
    
    # Body
    body = doc.add_paragraph("I am writing to formally request the reassignment of our current caseworker, ")
    body.add_run("{{ current_worker_name }}").italic = True
    if province_code:
        body.add_run(f", under the authority of the {legislation} in {province_code}.")
    else:
        body.add_run(", under the authority of the {{ legal_act }} in {{ province }}.")
    
    doc.add_paragraph("This request is made in good faith due to the following concerns and circumstances:")
    doc.add_paragraph()
    
    concerns = doc.add_paragraph("{{ description_of_concerns }}")
    concerns.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    doc.add_paragraph("I believe that continuing to work with the current assigned worker may compromise the integrity of our case and is not conducive to a respectful or constructive working relationship. It is in the best interest of all parties, especially the child(ren) involved, that a new worker be assigned who can approach this case with impartiality and professionalism.")
    doc.add_paragraph()
    
    doc.add_paragraph("Please confirm receipt of this request and provide written acknowledgment of next steps within 5 business days.")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("{{ name }}")
    doc.add_paragraph("{{ email }}")
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    
    # Save the document
    doc.save(template_path)
    print(f"Created template: {template_path}")
    
    return template_path

def create_all_provincial_templates():
    """Create CAS worker reassignment templates for all provinces"""
    provinces = ["ON", "BC", "AB", "QC", "MB", "SK", "NS", "NB", "NL", "PE", "NT", "NU", "YT"]
    
    # Create generic template
    create_cas_worker_reassignment_template()
    
    # Create province-specific templates
    count = 0
    for province in provinces:
        create_cas_worker_reassignment_template(province)
        count += 1
        
    # Create French Quebec template
    create_french_quebec_template()
    
    print(f"Created {count} provincial templates.")

# Create all templates if run directly
if __name__ == "__main__":
    create_all_provincial_templates()