#!/usr/bin/env python3
"""
Template Creation Script for CAS Records Request Documents
Creates templates for all Canadian provinces with the appropriate legal references
"""

import os
import sys
from docxtpl import DocxTemplate
from docx import Document
import docx

def get_cas_provincial_legislation(province_code):
    """Return the appropriate child protection legislation name based on province code"""
    legislation_map = {
        "ON": "Child, Youth and Family Services Act",
        "BC": "Child, Family and Community Service Act",
        "AB": "Child, Youth and Family Enhancement Act",
        "QC": "Youth Protection Act",
        "MB": "Child and Family Services Act",
        "SK": "Child and Family Services Act", 
        "NS": "Children and Family Services Act",
        "NB": "Family Services Act",
        "NL": "Children, Youth and Families Act",
        "PE": "Child Protection Act",
        "NT": "Child and Family Services Act",
        "NU": "Child and Family Services Act",
        "YT": "Child and Family Services Act"
    }
    
    return legislation_map.get(province_code, "Child Protection Legislation")

def get_privacy_legislation(province_code):
    """Return the appropriate privacy legislation based on province code"""
    privacy_map = {
        "ON": "Freedom of Information and Protection of Privacy Act (FIPPA) and the Personal Health Information Protection Act (PHIPA)",
        "BC": "Freedom of Information and Protection of Privacy Act (FOIPPA)",
        "AB": "Freedom of Information and Protection of Privacy Act (FOIP)",
        "QC": "Act Respecting Access to Documents Held by Public Bodies and the Protection of Personal Information",
        "MB": "Freedom of Information and Protection of Privacy Act (FIPPA) and The Personal Health Information Act (PHIA)",
        "SK": "Local Authority Freedom of Information and Protection of Privacy Act (LA FOIP)",
        "NS": "Freedom of Information and Protection of Privacy Act (FOIPOP)",
        "NB": "Right to Information and Protection of Privacy Act (RTIPPA)",
        "NL": "Access to Information and Protection of Privacy Act (ATIPPA)",
        "PE": "Freedom of Information and Protection of Privacy Act (FOIPP)",
        "NT": "Access to Information and Protection of Privacy Act (ATIPP)",
        "NU": "Access to Information and Protection of Privacy Act (ATIPP)",
        "YT": "Access to Information and Protection of Privacy Act (ATIPP)"
    }
    
    return privacy_map.get(province_code, "Freedom of Information and Protection of Privacy legislation")

def get_agency_name(province_code):
    """Return the appropriate agency name based on province code"""
    agency_map = {
        "ON": "Children's Aid Society",
        "BC": "Ministry of Children and Family Development",
        "AB": "Child and Family Services",
        "QC": "Direction de la protection de la jeunesse (DPJ)",
        "MB": "Child and Family Services",
        "SK": "Ministry of Social Services",
        "NS": "Department of Community Services",
        "NB": "Department of Social Development",
        "NL": "Department of Children, Seniors and Social Development",
        "PE": "Child Protection Services",
        "NT": "Child and Family Services",
        "NU": "Department of Family Services",
        "YT": "Family and Children's Services"
    }
    
    return agency_map.get(province_code, "Child Protection Services")

def create_french_quebec_template():
    """Create a specialized French template for Quebec DPJ"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    output_path = f"{template_dir}/cas_records_request_QC_FR.docx"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('DEMANDE D\'ACCÈS AU DOSSIER', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add DPJ info
    doc.add_paragraph("À:")
    doc.add_paragraph("Direction de la protection de la jeunesse (DPJ)")
    doc.add_paragraph("{{agency_address}}")
    doc.add_paragraph("{{agency_city}}, QC {{agency_postal_code}}")
    doc.add_paragraph()
    
    # Add parent info
    doc.add_paragraph("De:")
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("{{parent_address}}")
    doc.add_paragraph("{{parent_city}}, QC {{parent_postal_code}}")
    doc.add_paragraph()
    
    # Add reference information
    doc.add_paragraph("Référence: {{case_reference}}")
    doc.add_paragraph("Concernant: {{child_name}}, né(e) le {{child_dob}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph()
    
    doc.add_paragraph("Je vous écris pour demander formellement l'accès à tous les dossiers et documents concernant mon enfant, {{child_name}}, conformément à la Loi sur la protection de la jeunesse et à la Loi sur l'accès aux documents des organismes publics et sur la protection des renseignements personnels.")
    doc.add_paragraph()
    
    # Specific records requested
    doc.add_paragraph("Je demande spécifiquement les documents suivants:")
    doc.add_paragraph("1. Tous les rapports d'évaluation et d'enquête")
    doc.add_paragraph("2. Tous les plans d'intervention")
    doc.add_paragraph("3. Tous les rapports de visite et notes de suivi")
    doc.add_paragraph("4. Tous les rapports médicaux et psychologiques")
    doc.add_paragraph("5. Les transcriptions des entrevues avec mon enfant")
    doc.add_paragraph("6. Tous les rapports judiciaires et décisions du tribunal")
    doc.add_paragraph("7. Toute correspondance concernant mon enfant")
    doc.add_paragraph("8. {{additional_records}}")
    doc.add_paragraph()
    
    # Timeframe
    doc.add_paragraph("Période concernée: {{timeframe}}")
    doc.add_paragraph()
    
    # Legal reference
    doc.add_paragraph("Cette demande est faite conformément à mon droit d'accès aux informations en tant que parent, tel que défini par la loi. Je comprends que certaines informations peuvent être caviardées pour protéger la confidentialité des tiers, tel que prévu par la loi.")
    doc.add_paragraph()
    
    # Format request
    doc.add_paragraph("Je demande que ces documents me soient fournis sous forme de {{format_requested}} (copies papier ou électroniques).")
    doc.add_paragraph()
    
    # Authorization for release
    doc.add_paragraph("AUTORISATION:", style='Heading2')
    doc.add_paragraph("Je, {{parent_name}}, autorise la divulgation des documents mentionnés ci-dessus concernant {{child_name}}.")
    doc.add_paragraph()
    
    # Timeframe for response
    doc.add_paragraph("Conformément à la loi, je m'attends à recevoir une réponse à cette demande dans les 20 jours suivant sa réception.")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Je vous remercie de votre attention à cette demande. N'hésitez pas à me contacter si vous avez besoin d'informations supplémentaires pour traiter cette demande.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Veuillez agréer mes salutations distinguées,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("Téléphone: {{parent_phone}}")
    doc.add_paragraph("Courriel: {{parent_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_cas_records_request_template(province_code=None):
    """Create a specialized template for CAS records request"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    # Create a base template name
    base_template_path = f"{template_dir}/cas_records_request.docx"
    
    # If a province code is provided, create a province-specific template
    if province_code:
        output_path = f"{template_dir}/cas_records_request_{province_code}.docx"
        child_legislation = get_cas_provincial_legislation(province_code)
        privacy_legislation = get_privacy_legislation(province_code)
        agency_name = get_agency_name(province_code)
    else:
        output_path = base_template_path
        child_legislation = "Child Protection Act"
        privacy_legislation = "Freedom of Information and Protection of Privacy legislation"
        agency_name = "Child Protection Services"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('REQUEST FOR ACCESS TO RECORDS', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add agency info
    doc.add_paragraph("To:")
    doc.add_paragraph(f"{agency_name}")
    doc.add_paragraph("{{agency_address}}")
    doc.add_paragraph("{{agency_city}}, {{agency_province}} {{agency_postal_code}}")
    doc.add_paragraph()
    
    # Add parent info
    doc.add_paragraph("From:")
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("{{parent_address}}")
    doc.add_paragraph("{{parent_city}}, {{parent_province}} {{parent_postal_code}}")
    doc.add_paragraph()
    
    # Add reference information
    doc.add_paragraph("Reference: {{case_reference}}")
    doc.add_paragraph("Regarding: {{child_name}}, born {{child_dob}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph("Dear Sir/Madam,")
    doc.add_paragraph()
    
    doc.add_paragraph(f"I am writing to formally request access to all records and documents pertaining to my child, {{child_name}}, in accordance with the {child_legislation} and the {privacy_legislation}.")
    doc.add_paragraph()
    
    # Specific records requested
    doc.add_paragraph("I am specifically requesting the following records:")
    doc.add_paragraph("1. All assessment and investigation reports")
    doc.add_paragraph("2. All service plans or plans of care")
    doc.add_paragraph("3. All visit reports and case notes")
    doc.add_paragraph("4. All medical and psychological reports")
    doc.add_paragraph("5. Transcripts of interviews with my child")
    doc.add_paragraph("6. All court reports and judicial decisions")
    doc.add_paragraph("7. Any correspondence regarding my child")
    doc.add_paragraph("8. {{additional_records}}")
    doc.add_paragraph()
    
    # Timeframe
    doc.add_paragraph("Time period covered: {{timeframe}}")
    doc.add_paragraph()
    
    # Legal reference
    doc.add_paragraph(f"This request is made pursuant to my right of access to information as a parent, as defined by law. I understand that certain information may be redacted to protect third-party confidentiality, as provided for in the {privacy_legislation}.")
    doc.add_paragraph()
    
    # Format request
    doc.add_paragraph("I request that these records be provided to me in {{format_requested}} form (paper copies or electronic format).")
    doc.add_paragraph()
    
    # Authorization for release
    doc.add_paragraph("AUTHORIZATION:", style='Heading2')
    doc.add_paragraph("I, {{parent_name}}, authorize the release of the above-mentioned records concerning {{child_name}}.")
    doc.add_paragraph()
    
    # Timeframe for response
    doc.add_paragraph("In accordance with the statutory requirements, I expect a response to this request within 30 days of receipt.")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Thank you for your attention to this request. Please do not hesitate to contact me if you require any additional information to process this request.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("Phone: {{parent_phone}}")
    doc.add_paragraph("Email: {{parent_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_all_provincial_templates():
    """Create CAS records request templates for all provinces"""
    provinces = ["ON", "BC", "AB", "QC", "MB", "SK", "NS", "NB", "NL", "PE", "NT", "NU", "YT"]
    template_count = 0
    
    # First create generic template
    create_cas_records_request_template()
    
    # Then create province-specific templates
    for province in provinces:
        create_cas_records_request_template(province)
        template_count += 1
    
    # Create Quebec French template
    create_french_quebec_template()
    
    print(f"Created {template_count} provincial CAS records request templates.")

if __name__ == "__main__":
    create_all_provincial_templates()