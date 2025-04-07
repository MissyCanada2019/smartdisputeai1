#!/usr/bin/env python3
"""
Template Creation Script for CAS Decision Appeal Documents
Creates templates for all Canadian provinces with the appropriate legal references
"""

import os
import sys
from docxtpl import DocxTemplate
from docx import Document
import docx

def get_cas_provincial_legislation(province_code):
    """Return the appropriate child protection legislation name based on province code"""
    legislation_map = {
        "ON": "Child, Youth and Family Services Act",
        "BC": "Child, Family and Community Service Act",
        "AB": "Child, Youth and Family Enhancement Act",
        "QC": "Youth Protection Act",
        "MB": "Child and Family Services Act",
        "SK": "Child and Family Services Act",
        "NS": "Children and Family Services Act",
        "NB": "Family Services Act",
        "NL": "Children, Youth and Families Act",
        "PE": "Child Protection Act",
        "NT": "Child and Family Services Act",
        "NU": "Child and Family Services Act",
        "YT": "Child and Family Services Act"
    }
    
    return legislation_map.get(province_code, "Child Protection Legislation")

def get_appeal_authority(province_code):
    """Return the appropriate appeal authority based on province code"""
    authority_map = {
        "ON": "Child and Family Services Review Board",
        "BC": "Ministry of Children and Family Development Complaints Process and the BC Ombudsperson",
        "AB": "Appeals Secretariat",
        "QC": "Administrative Tribunal of Quebec, Social Affairs Division",
        "MB": "Social Services Appeal Board",
        "SK": "Social Services Appeal Board",
        "NS": "Minister of Community Services",
        "NB": "Family Services Appeal Board",
        "NL": "Children, Youth and Families Appeal Board",
        "PE": "Child Protection Act Appeal Panel",
        "NT": "Director of Child and Family Services",
        "NU": "Director of Child and Family Services",
        "YT": "Director of Family and Children's Services"
    }
    
    return authority_map.get(province_code, "Child Protection Appeal Authority")

def get_appeal_timeframe(province_code):
    """Return the appropriate appeal timeframe based on province code"""
    timeframe_map = {
        "ON": "30 days from the date of the decision",
        "BC": "30 days from the date of the decision",
        "AB": "30 days from the date of the decision",
        "QC": "60 days from the date of the decision",
        "MB": "30 days from the date of the decision",
        "SK": "30 days from the date of the decision",
        "NS": "30 days from the date of the decision",
        "NB": "30 days from the date of the decision",
        "NL": "10 days from the date of the decision",
        "PE": "20 days from the date of the decision",
        "NT": "10 days from the date of the decision",
        "NU": "10 days from the date of the decision",
        "YT": "30 days from the date of the decision"
    }
    
    return timeframe_map.get(province_code, "30 days from the date of the decision")

def get_agency_name(province_code):
    """Return the appropriate agency name based on province code"""
    agency_map = {
        "ON": "Children's Aid Society",
        "BC": "Ministry of Children and Family Development",
        "AB": "Child and Family Services",
        "QC": "Direction de la protection de la jeunesse (DPJ)",
        "MB": "Child and Family Services",
        "SK": "Ministry of Social Services",
        "NS": "Department of Community Services",
        "NB": "Department of Social Development",
        "NL": "Department of Children, Seniors and Social Development",
        "PE": "Child Protection Services",
        "NT": "Child and Family Services",
        "NU": "Department of Family Services",
        "YT": "Family and Children's Services"
    }
    
    return agency_map.get(province_code, "Child Protection Services")

def create_french_quebec_template():
    """Create a specialized French template for Quebec DPJ"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    output_path = f"{template_dir}/cas_appeal_QC_FR.docx"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('DEMANDE D\'APPEL D\'UNE DÉCISION DE LA PROTECTION DE LA JEUNESSE', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add appeal authority info
    doc.add_paragraph("À:")
    doc.add_paragraph("Tribunal administratif du Québec, Section des affaires sociales")
    doc.add_paragraph("{{appeal_authority_address}}")
    doc.add_paragraph("{{appeal_authority_city}}, QC {{appeal_authority_postal_code}}")
    doc.add_paragraph()
    
    # Add copy to DPJ
    doc.add_paragraph("Copie à:")
    doc.add_paragraph("Direction de la protection de la jeunesse (DPJ)")
    doc.add_paragraph("{{agency_address}}")
    doc.add_paragraph("{{agency_city}}, QC {{agency_postal_code}}")
    doc.add_paragraph()
    
    # Add parent info
    doc.add_paragraph("De:")
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("{{parent_address}}")
    doc.add_paragraph("{{parent_city}}, QC {{parent_postal_code}}")
    doc.add_paragraph()
    
    # Add reference information
    doc.add_paragraph("Référence: {{case_reference}}")
    doc.add_paragraph("Concernant: {{child_name}}, né(e) le {{child_dob}}")
    doc.add_paragraph("Date de la décision contestée: {{decision_date}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph()
    
    doc.add_paragraph("Par la présente, je souhaite formellement faire appel de la décision rendue par la Direction de la protection de la jeunesse (DPJ) le {{decision_date}} concernant mon enfant {{child_name}}, conformément à la Loi sur la protection de la jeunesse et au droit d'appel prévu à la Section III de cette loi.")
    doc.add_paragraph()
    
    # Decision being appealed
    doc.add_paragraph("DÉCISION CONTESTÉE:", style='Heading2')
    doc.add_paragraph("{{decision_details}}")
    doc.add_paragraph()
    
    # Grounds for appeal
    doc.add_paragraph("MOTIFS DE L'APPEL:", style='Heading2')
    doc.add_paragraph("Je conteste cette décision pour les motifs suivants:")
    doc.add_paragraph("{{appeal_grounds}}")
    doc.add_paragraph()
    
    # Legal basis
    doc.add_paragraph("FONDEMENT JURIDIQUE:", style='Heading2')
    doc.add_paragraph("Cet appel est fondé sur les dispositions suivantes de la Loi sur la protection de la jeunesse:")
    doc.add_paragraph("{{legal_basis}}")
    doc.add_paragraph()
    
    # Supporting evidence
    doc.add_paragraph("PREUVE À L'APPUI:", style='Heading2')
    doc.add_paragraph("Je joins les documents suivants à l'appui de ma demande:")
    doc.add_paragraph("{{supporting_evidence}}")
    doc.add_paragraph()
    
    # Requested outcome
    doc.add_paragraph("DÉCISION DEMANDÉE:", style='Heading2')
    doc.add_paragraph("Je demande respectueusement que:")
    doc.add_paragraph("{{requested_outcome}}")
    doc.add_paragraph()
    
    # Timeframe reference
    doc.add_paragraph("Je comprends que le délai légal pour déposer cet appel est de 60 jours à compter de la date de la décision, et cette demande est soumise dans ce délai.")
    doc.add_paragraph()
    
    # Request for hearing
    doc.add_paragraph("Je demande une audience pour pouvoir présenter mes arguments en personne. Je suis disponible aux dates suivantes:")
    doc.add_paragraph("{{available_dates}}")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Je vous remercie de votre attention à cette demande d'appel. Je suis disponible pour fournir tout renseignement supplémentaire qui pourrait être requis.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Veuillez agréer mes salutations distinguées,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("Téléphone: {{parent_phone}}")
    doc.add_paragraph("Courriel: {{parent_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_cas_appeal_template(province_code=None):
    """Create a specialized template for CAS decision appeal"""
    template_dir = "templates/docs"
    os.makedirs(template_dir, exist_ok=True)
    
    # Create a base template name
    base_template_path = f"{template_dir}/cas_appeal.docx"
    
    # If a province code is provided, create a province-specific template
    if province_code:
        output_path = f"{template_dir}/cas_appeal_{province_code}.docx"
        legislation = get_cas_provincial_legislation(province_code)
        appeal_authority = get_appeal_authority(province_code)
        appeal_timeframe = get_appeal_timeframe(province_code)
        agency_name = get_agency_name(province_code)
    else:
        output_path = base_template_path
        legislation = "Child Protection Act"
        appeal_authority = "Child Protection Appeal Authority"
        appeal_timeframe = "30 days from the date of the decision"
        agency_name = "Child Protection Services"
    
    # Create new document
    doc = Document()
    
    # Add header
    header = doc.add_heading('APPEAL OF CHILD PROTECTION DECISION', level=1)
    header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    doc.add_paragraph("Date: {{current_date}}")
    
    # Add appeal authority info
    doc.add_paragraph("To:")
    doc.add_paragraph(f"{appeal_authority}")
    doc.add_paragraph("{{appeal_authority_address}}")
    doc.add_paragraph("{{appeal_authority_city}}, {{appeal_authority_province}} {{appeal_authority_postal_code}}")
    doc.add_paragraph()
    
    # Add copy to agency
    doc.add_paragraph("Copy to:")
    doc.add_paragraph(f"{agency_name}")
    doc.add_paragraph("{{agency_address}}")
    doc.add_paragraph("{{agency_city}}, {{agency_province}} {{agency_postal_code}}")
    doc.add_paragraph()
    
    # Add parent info
    doc.add_paragraph("From:")
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("{{parent_address}}")
    doc.add_paragraph("{{parent_city}}, {{parent_province}} {{parent_postal_code}}")
    doc.add_paragraph()
    
    # Add reference information
    doc.add_paragraph("Reference: {{case_reference}}")
    doc.add_paragraph("Regarding: {{child_name}}, born {{child_dob}}")
    doc.add_paragraph("Date of decision being appealed: {{decision_date}}")
    doc.add_paragraph()
    
    # Add main content
    doc.add_paragraph("Dear Sir/Madam,")
    doc.add_paragraph()
    
    doc.add_paragraph(f"I am writing to formally appeal the decision made by the {agency_name} on {{decision_date}} regarding my child {{child_name}}, pursuant to the {legislation} and the appeal rights provided therein.")
    doc.add_paragraph()
    
    # Decision being appealed
    doc.add_paragraph("DECISION BEING APPEALED:", style='Heading2')
    doc.add_paragraph("{{decision_details}}")
    doc.add_paragraph()
    
    # Grounds for appeal
    doc.add_paragraph("GROUNDS FOR APPEAL:", style='Heading2')
    doc.add_paragraph("I am appealing this decision on the following grounds:")
    doc.add_paragraph("{{appeal_grounds}}")
    doc.add_paragraph()
    
    # Legal basis
    doc.add_paragraph("LEGAL BASIS:", style='Heading2')
    doc.add_paragraph(f"This appeal is based on the following provisions of the {legislation}:")
    doc.add_paragraph("{{legal_basis}}")
    doc.add_paragraph()
    
    # Supporting evidence
    doc.add_paragraph("SUPPORTING EVIDENCE:", style='Heading2')
    doc.add_paragraph("I am attaching the following documents in support of this appeal:")
    doc.add_paragraph("{{supporting_evidence}}")
    doc.add_paragraph()
    
    # Requested outcome
    doc.add_paragraph("REQUESTED OUTCOME:", style='Heading2')
    doc.add_paragraph("I respectfully request that:")
    doc.add_paragraph("{{requested_outcome}}")
    doc.add_paragraph()
    
    # Timeframe reference
    doc.add_paragraph(f"I understand that the statutory timeframe for filing this appeal is {appeal_timeframe}, and this appeal is being submitted within that timeframe.")
    doc.add_paragraph()
    
    # Request for hearing
    doc.add_paragraph("I request a hearing to present my arguments in person. I am available on the following dates:")
    doc.add_paragraph("{{available_dates}}")
    doc.add_paragraph()
    
    # Closing
    doc.add_paragraph("Thank you for your attention to this appeal. I am available to provide any additional information that may be required.")
    doc.add_paragraph()
    
    # Add signature
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("{{parent_name}}")
    doc.add_paragraph("Phone: {{parent_phone}}")
    doc.add_paragraph("Email: {{parent_email}}")
    
    # Save the document
    doc.save(output_path)
    print(f"Created template: {output_path}")
    
    return output_path

def create_all_provincial_templates():
    """Create CAS appeal templates for all provinces"""
    provinces = ["ON", "BC", "AB", "QC", "MB", "SK", "NS", "NB", "NL", "PE", "NT", "NU", "YT"]
    template_count = 0
    
    # First create generic template
    create_cas_appeal_template()
    
    # Then create province-specific templates
    for province in provinces:
        create_cas_appeal_template(province)
        template_count += 1
    
    # Create Quebec French template
    create_french_quebec_template()
    
    print(f"Created {template_count} provincial CAS appeal templates.")

if __name__ == "__main__":
    create_all_provincial_templates()